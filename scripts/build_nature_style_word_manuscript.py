from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
TABLES = OUT / "tables"
FIGS = OUT / "figures"
MANUSCRIPT_DIR = OUT / "manuscript_text"
DOCX_OUT = MANUSCRIPT_DIR / "state_capacity_nature_style_manuscript.docx"

PRESET = {
    "name": "narrative_proposal",
    "page_width": 8.5,
    "page_height": 11.0,
    "margin": 1.0,
    "header_footer": 0.492,
    "content_width_dxa": 9360,
    "table_indent_dxa": 120,
    "cell_margins": {"top": 80, "bottom": 80, "start": 120, "end": 120},
    "body_font": "Calibri",
    "body_size": 11,
    "body_after": 8,
    "body_line": 1.333,
    "heading_blue": "2E74B5",
    "heading_dark": "1F4D78",
    "table_header_fill": "F4F6F9",
}


FIGURE_CAPTIONS = {
    "figure_1_pipeline_and_claim_coverage.png": (
        "Figure 1 | Study pipeline and claim coverage. Dataset scale and master "
        "claim-audit coverage across state, capacity, interaction, load and "
        "control constructs."
    ),
    "figure_2_ann_gate.png": (
        "Figure 2 | Artificial-agent intervention gate. Matched state-versus-"
        "capacity classification succeeds under residualized fingerprints, but "
        "residualized state-severity and hybrid-state recovery fail the stricter gate."
    ),
    "figure_3_ds007554_discovery.png": (
        "Figure 3 | ds007554 reconstructed discovery analysis. Leave-one-participant-"
        "out model comparison and bootstrap state/capacity associations using "
        "push-button reconstructed correctness."
    ),
    "figure_4_recurrent_dynamics.png": (
        "Figure 4 | Recurrent dynamics. Associations between state/capacity "
        "profiles and fitted GRU trajectory geometry, plus latent decoder performance."
    ),
    "figure_5_ds007554_neurophysiology.png": (
        "Figure 5 | ds007554 neurophysiology. Direct state/capacity associations "
        "with EEG, ECG, fNIRS and physiology-derived features after coordinate repair."
    ),
    "figure_6_cog_bci_validation.png": (
        "Figure 6 | COG-BCI validation. State and capacity associations with "
        "COG-BCI behavior and EEG features, including strong state associations "
        "with RT variability and lapse."
    ),
    "figure_7_tu_hbn_validation.png": (
        "Figure 7 | TU Berlin and HBN validation. TU Berlin load/capacity-pressure "
        "effects and HBN scalability/geometry validation."
    ),
    "figure_8_robustness_falsification.png": (
        "Figure 8 | Robustness and falsification. Hidden-size sensitivity, "
        "residualized tests, leave-one-dataset/task checks, random/shuffled controls, "
        "permutation controls and baseline challenges."
    ),
}


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:], 16))


def fmt_num(value, sig: int = 3) -> str:
    if value is None or pd.isna(value):
        return ""
    try:
        x = float(value)
    except (TypeError, ValueError):
        return str(value)
    if x == 0:
        return "0"
    if abs(x) < 0.001 or abs(x) >= 10000:
        return f"{x:.{sig}e}"
    if abs(x) < 1:
        return f"{x:.{sig}g}"
    return f"{x:.{sig}f}".rstrip("0").rstrip(".")


def fmt_p(value) -> str:
    if value is None or pd.isna(value) or str(value) == "":
        return ""
    x = float(value)
    if x < 0.001:
        return f"{x:.2e}"
    return f"{x:.3f}".rstrip("0").rstrip(".")


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.name = PRESET["body_font"]
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(PRESET["table_indent_dxa"]))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in PRESET["cell_margins"].items():
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(val))
                node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "B8C2CC", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(
    doc: Document,
    title: str,
    source: str,
    columns: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_dxa: Sequence[int],
) -> None:
    p = doc.add_paragraph()
    p.style = "Table Caption"
    r = p.add_run(title)
    r.bold = True
    p.add_run(f" Source data: {source}.")

    row_list = list(rows)
    table = doc.add_table(rows=len(row_list) + 1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)

    for idx, col in enumerate(columns):
        set_cell_text(table.rows[0].cells[idx], col, bold=True, size=8.2)
        set_cell_shading(table.rows[0].cells[idx], PRESET["table_header_fill"])
    for r_idx, row in enumerate(row_list, 1):
        for c_idx, value in enumerate(row):
            set_cell_text(table.rows[r_idx].cells[c_idx], value, size=8.0)

    doc.add_paragraph("")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style is None else p.alignment
    p.add_run(text)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(9)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PRESET["page_width"])
    section.page_height = Inches(PRESET["page_height"])
    section.top_margin = Inches(PRESET["margin"])
    section.bottom_margin = Inches(PRESET["margin"])
    section.left_margin = Inches(PRESET["margin"])
    section.right_margin = Inches(PRESET["margin"])
    section.header_distance = Inches(PRESET["header_footer"])
    section.footer_distance = Inches(PRESET["header_footer"])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRESET["body_font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal.font.color.rgb = rgb("111111")
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["body_line"]

    for name, size, color, before, after in [
        ("Heading 1", 16, PRESET["heading_blue"], 18, 10),
        ("Heading 2", 13, PRESET["heading_blue"], 12, 6),
        ("Heading 3", 12, PRESET["heading_dark"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = PRESET["body_font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = PRESET["body_font"]
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = rgb("333333")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.05

    table_caption = styles.add_style("Table Caption", 1)
    table_caption.font.name = PRESET["body_font"]
    table_caption.font.size = Pt(9)
    table_caption.font.color.rgb = rgb("333333")
    table_caption.paragraph_format.space_before = Pt(4)
    table_caption.paragraph_format.space_after = Pt(4)
    table_caption.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = "State and capacity in compact recurrent models"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = rgb("666666")


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "State and capacity in compact recurrent models of human task behavior: "
        "a multi-dataset validation and falsification study"
    )
    run.font.name = PRESET["body_font"]
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = rgb("0B2545")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta_run = meta.add_run(
        "Nature-style manuscript draft generated from the reconstructed "
        "state_capacity_tinyrnn pipeline. All numerical claims are taken from "
        "scripted output tables generated on 2026-06-10."
    )
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = rgb("555555")

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Behavioral impairment can arise because a system is temporarily in a "
        "poor operating state, because it has limited representational capacity, "
        "or because both mechanisms interact. These alternatives are often "
        "confounded by overall accuracy. Here we reconstructed a state-capacity "
        "analysis pipeline around compact gated recurrent neural networks trained "
        "on directly downloadable human behavioral datasets with EEG, fNIRS and "
        "ECG extensions. Artificial-agent perturbations established that "
        "state-like and capacity-like manipulations can be classified under "
        "matched performance, but residualized recovery of a single state-severity "
        "axis failed. We therefore treated state and capacity as multidimensional "
        "profiles rather than clean scalar traits. Across COG-BCI, ds007554, TU "
        "Berlin and HBN Release 4, the strongest evidence supported capacity as "
        "a geometry- and pressure-related profile. State showed meaningful but "
        "less stable evidence, especially in within-person behavioral reliability, "
        "whereas direct state-capacity interaction evidence remained mixed. "
        "These results support a cautious reframing: compact recurrent models can "
        "separate useful state-like reliability profiles from capacity-like "
        "geometry profiles, but current evidence justifies stronger claims for "
        "capacity pressure than for state as an independently recovered coordinate.",
    )


def add_main_text(doc: Document) -> None:
    add_heading(doc, "Main", 1)
    for text in [
        "Momentary task failure is ambiguous. A person may respond poorly because "
        "the same cognitive system is operating in an unfavorable state, because "
        "the system lacks sufficient representational capacity for the task, or "
        "because task demand exposes a capacity limit only under certain states. "
        "Accuracy alone cannot distinguish these alternatives.",
        "Compact recurrent neural networks provide a practical way to formalize "
        "this distinction. Operating perturbations such as lapse, noise, gain or "
        "memory decay can mimic state changes without changing structural resources; "
        "changes in hidden size or recurrent resources can mimic capacity limits. "
        "If these manipulations leave separable behavioral and dynamical "
        "fingerprints after performance matching, they provide an operational basis "
        "for projecting human behavior into state-like and capacity-like profiles.",
        "The guiding premise was deliberately conservative. State and capacity were "
        "not assumed to be single latent variables. Each was represented as a "
        "multidimensional profile, and claims were graded by convergence across "
        "artificial-agent recovery, human behavior, recurrent dynamics, physiology, "
        "external datasets and negative controls.",
    ]:
        add_para(doc, text)


def add_results(doc: Document) -> None:
    add_heading(doc, "Results", 1)
    add_heading(doc, "A multi-dataset event table anchors the reconstruction", 2)
    add_para(
        doc,
        "The reconstructed event table joined four directly downloadable datasets. "
        "COG-BCI contributed 55,854 included events from 29 participants, ds007554 "
        "contributed 31,053 events from 30 participants with 10,896 reconstructed "
        "supervised events, TU Berlin contributed 22,096 included events from 26 "
        "participants, and HBN Release 4 contributed 180,824 included events from "
        "324 participants. Across the full training set, Step 08 produced 1,283 "
        "session-task state rows and 390 participant capacity rows.",
    )

    counts = pd.read_csv(TABLES / "event_counts_by_dataset.csv")
    add_table(
        doc,
        "Table 1 | Unified event-table scale.",
        "event_counts_by_dataset.csv",
        ["Dataset", "Included events", "Subjects", "Sessions", "Tasks", "Events with correctness"],
        [
            [
                r["dataset"],
                fmt_num(r["n_included_events"], 4),
                fmt_num(r["n_subjects"], 4),
                fmt_num(r["n_sessions"], 4),
                fmt_num(r["n_tasks"], 4),
                fmt_num(r["n_events_with_correct"], 4),
            ]
            for _, r in counts.iterrows()
        ],
        [2200, 1500, 1100, 1100, 900, 2560],
    )

    add_heading(doc, "Artificial perturbations separate classes but not a scalar state axis", 2)
    add_para(
        doc,
        "Artificial agents showed that state-like and capacity-like perturbations "
        "can be distinguished under matched performance. The residualized "
        "state-versus-capacity classifier remained above shuffled controls across "
        "task families, with overall balanced accuracy of 0.819. However, the "
        "stricter recovery criterion failed for state: residualized state-severity "
        "recovery was only rho = 0.364, and the hybrid state axis failed permutation "
        "testing. This failure defines the main claim boundary for the study: human "
        "state findings are interpreted as reliability-profile evidence unless "
        "independently validated.",
    )

    add_heading(doc, "Human projection improves prediction, but descriptive behavior remains a hard baseline", 2)
    add_para(
        doc,
        "In ds007554, reconstructed push-button correctness made supervised discovery "
        "eligible. The additive state-capacity model improved over the task/dataset "
        "baseline, but the behavioral descriptive baseline remained much stronger. "
        "The machine-projection model had LOPO RMSE = 0.0292, the additive "
        "state-capacity model had LOPO RMSE = 0.0446, and the behavioral descriptive "
        "baseline had LOPO RMSE = 0.0093. In the pooled supervised context, additive "
        "state-capacity prediction improved over the task/dataset baseline, and a "
        "small interaction term further improved RMSE.",
    )

    model_cmp = pd.read_csv(TABLES / "ds007554_discovery_model_comparison.csv")
    selected = model_cmp[
        model_cmp["model_name"].isin(
            [
                "task_dataset",
                "additive_state_capacity",
                "state_capacity_interaction",
                "machine_projection_additive",
                "behavioral_descriptive",
                "random_axis_control",
                "shuffled_coordinate_control",
            ]
        )
    ]
    add_table(
        doc,
        "Table 2 | Leave-one-participant-out prediction model comparison.",
        "ds007554_discovery_model_comparison.csv",
        ["Scope", "Model", "Role", "n rows", "Participants", "LOPO RMSE"],
        [
            [
                r["analysis_scope"].replace("_", " "),
                r["model_name"].replace("_", " "),
                r["claim_role"].replace("_", " "),
                fmt_num(r["n_rows"], 4),
                fmt_num(r["n_participants"], 4),
                fmt_num(r["lopo_rmse"], 4),
            ]
            for _, r in selected.iterrows()
        ],
        [2350, 2150, 1750, 850, 950, 1310],
    )

    add_heading(doc, "Capacity has the clearest recurrent-dynamics signature", 2)
    add_para(
        doc,
        "Capacity-like profiles aligned strongly with fitted recurrent geometry. "
        "Across 1,283 trajectories and 115,979 hidden-state rows, capacity was "
        "associated with trajectory covariance rank and related geometry measures. "
        "The strongest reported association was trajectory covariance rank with "
        "capacity_parameter_resource_z (rho = 0.713, p = 1.53e-199). State-related "
        "profiles were also associated with several geometry summaries, but the "
        "interpretation is more fragile because the artificial-agent residualized "
        "state gate failed.",
    )

    dynamics = pd.read_csv(TABLES / "recurrent_dynamics_state_capacity_tests.csv")
    dyn_sel = dynamics[
        ((dynamics["outcome"] == "trajectory_cov_rank") & dynamics["predictor"].isin(["capacity_parameter_resource_z", "optimized_state_profile_z"]))
        | ((dynamics["outcome"] == "trajectory_radius") & dynamics["predictor"].isin(["capacity_parameter_resource_z", "optimized_state_profile_z"]))
        | ((dynamics["outcome"] == "hidden_variability") & dynamics["predictor"].isin(["capacity_parameter_resource_z", "optimized_state_profile_z"]))
    ]
    add_table(
        doc,
        "Table 3 | Recurrent-geometry associations.",
        "recurrent_dynamics_state_capacity_tests.csv",
        ["Outcome", "Predictor", "n", "Spearman rho", "p value"],
        [
            [
                r["outcome"].replace("_", " "),
                r["predictor"].replace("_", " "),
                fmt_num(r["n"], 4),
                fmt_num(r["spearman_rho"], 3),
                fmt_p(r["p_value"]),
            ]
            for _, r in dyn_sel.iterrows()
        ],
        [2450, 3200, 700, 1300, 1710],
    )

    add_heading(doc, "External validation separates transient state from capacity pressure", 2)
    add_para(
        doc,
        "COG-BCI provided the strongest state validation. State varied more within "
        "person than between person and predicted behavioral reliability: state "
        "predicted RT coefficient of variation (q = 1.50e-24), RT IQR "
        "(q = 1.58e-10), accuracy (q = 1.16e-07) and lapse rate (q = 1.16e-07). "
        "Capacity showed qualified support through cross-task consistency and EEG "
        "associations. TU Berlin validated capacity pressure: increasing N-back "
        "load reduced accuracy and slowed RT, and load-by-capacity interactions "
        "predicted accuracy (beta = 0.149, q = 4.69e-29) and RT (q = 0.0003). "
        "HBN Release 4 confirmed scalability and capacity-geometry alignment but "
        "did not provide strong participant-level state validation.",
    )

    validation_rows = [
        ["COG-BCI", "state -> RT CV", "state", "261", "0.291", "1.50e-24", "strong behavioral reliability"],
        ["COG-BCI", "state -> accuracy", "state", "261", "-0.127", "1.16e-07", "strong but signed as impairment axis"],
        ["COG-BCI", "capacity -> cross-task consistency", "capacity", "29", "0.502", "0.0117", "moderate external validation"],
        ["TU Berlin", "load x capacity -> accuracy", "interaction", "234", "0.149", "4.69e-29", "strong capacity pressure"],
        ["TU Berlin", "state instability -> lapse", "state", "234", "0.047", "ns", "negative after correction"],
        ["HBN", "capacity -> trajectory radius", "capacity", "306", "0.729", "4.17e-51", "strong scalable geometry"],
        ["HBN", "state mean -> accuracy", "state", "306", "0.045", "ns", "negative participant-level test"],
    ]
    add_table(
        doc,
        "Table 4 | Cross-dataset validation summary.",
        "cog_bci_validation_models.csv; tu_berlin_load_validation.csv; hbn_scalability_tests.csv",
        ["Dataset", "Test", "Construct", "n", "Estimate", "q/p", "Interpretation"],
        validation_rows,
        [1200, 2300, 1050, 650, 900, 1050, 2210],
    )

    add_heading(doc, "Neurophysiology is informative but still bounded", 2)
    add_para(
        doc,
        "Step 12 extracted ds007554 EEG, fNIRS and ECG features and attached repaired "
        "state/capacity coordinates. After push-button reconstruction, direct "
        "state/capacity physiology models became eligible. State showed exploratory "
        "associations with EEG spectral entropy and band-power features, while "
        "capacity showed stronger ECG/fNIRS associations in selected models. These "
        "analyses provide external alignment evidence, but they should not be "
        "written as direct neural proof that state and capacity coordinates are "
        "neural variables because ds007554 behavioral correctness was reconstructed "
        "from timing signals rather than distributed as explicit trial-correctness labels.",
    )

    add_heading(doc, "Robustness and falsification constrain the final claim", 2)
    add_para(
        doc,
        "Robustness analyses produced 160 robustness rows and 18 falsification rows. "
        "Capacity evidence was robust when tied to recurrent geometry and load/"
        "capacity-pressure validation. State remained limited by the failed ANN "
        "residualized state gate, nonsignificant TU Berlin state effects and weak HBN "
        "participant-level state tests. Random-axis and shuffled-coordinate controls "
        "did not reproduce the main coordinate gains, but the ds007554 behavioral "
        "descriptive baseline outperformed coordinate models. The master effects "
        "table contains 454 effect rows, including 279 state rows, 114 capacity rows "
        "and 32 interaction rows, intentionally including negative and failed-validation findings.",
    )

    summary = pd.read_csv(TABLES / "master_effects_summary.csv")
    add_table(
        doc,
        "Table 5 | Master claim-audit effect counts.",
        "master_effects_summary.csv",
        ["Construct", "Claim strength", "Number of effects"],
        [
            [r["construct"], r["claim_strength"], fmt_num(r["n_effects"], 4)]
            for _, r in summary.iterrows()
        ],
        [2600, 2600, 4160],
    )


def add_discussion(doc: Document) -> None:
    add_heading(doc, "Discussion", 1)
    for text in [
        "This study supports a cautious distinction between state-like reliability "
        "profiles and capacity-like recurrent-geometry profiles. The strongest "
        "evidence concerns capacity: capacity profiles aligned with recurrent "
        "geometry, generalized across HBN at scale, predicted COG-BCI cross-task "
        "consistency and moderated TU Berlin load-related decline. This makes "
        "capacity the more mature construct in the present analysis.",
        "State was meaningful but weaker. COG-BCI showed clear state associations "
        "with RT variability, lapse and accuracy, and ds007554 showed exploratory "
        "state-neurophysiology alignment. However, the artificial-agent state "
        "recovery gate failed after performance residualization, TU Berlin did not "
        "validate state effects after correction, and HBN participant-level state "
        "tests were weak. The correct interpretation is therefore not that state is "
        "absent, but that current state estimates are best described as exploratory "
        "behavioral reliability profiles rather than a validated scalar coordinate.",
        "Interaction evidence is similarly mixed. The direct ds007554 state-capacity "
        "interaction did not robustly improve over the additive model, whereas "
        "pooled interaction tests and TU Berlin load-by-capacity pressure were "
        "stronger. Thus, the safest interaction claim is about capacity pressure "
        "under increasing task demand, not a general state-by-capacity law.",
        "The study also shows the value of falsification. Random-axis and shuffled-"
        "coordinate controls did not explain the main coordinate effects, but simple "
        "descriptive behavioral summaries outperformed coordinate models in ds007554. "
        "This prevents overclaiming: compact recurrent coordinates are useful "
        "explanatory profiles, not automatically superior predictive models.",
    ]:
        add_para(doc, text)


def add_methods(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Methods", 1)
    methods = [
        (
            "Study design",
            "The study was implemented as a scripted reconstruction under "
            "state_capacity_tinyrnn. Each analysis step wrote tables, audit files, "
            "figures and source data. No empirical result was accepted into the "
            "manuscript unless it was produced by a script in the reconstructed "
            "folder and represented in the master claim audit.",
        ),
        (
            "Datasets",
            "The main datasets were OpenNeuro ds007554, COG-BCI, TU Berlin "
            "simultaneous EEG-NIRS and HBN Release 4. Datasets requiring manual "
            "login or uncontrolled access were excluded from confirmatory claims. "
            "A unified event schema represented dataset, subject, session, task, "
            "block, trial index, timestamp, condition, load, target, response, "
            "correctness, RT and inclusion flags.",
        ),
        (
            "Artificial agents",
            "Tiny recurrent agents were perturbed along state-like dimensions, "
            "including lapse, noise, gain, memory decay and temporal parameters, "
            "and capacity-like dimensions, including hidden size and recurrent "
            "resources. State and capacity perturbations were performance-matched "
            "before classification and recovery tests. Hybrid agents crossed state "
            "and capacity perturbations.",
        ),
        (
            "Human recurrent models",
            "Human behavior was modeled with compact GRUs using hidden sizes 1, 2, "
            "3, 4, 6 and 8. Splits included participant-level, session-blocked and "
            "odd/even mini-block validation. State profiles were session-task level "
            "and estimated from calibration windows while held-out events were not "
            "used for state estimation. Capacity profiles were participant level and "
            "summarized hidden-size selection, selection confidence, complexity "
            "preference, high-capacity advantage, load robustness and cross-task consistency.",
        ),
        (
            "State and capacity profiles",
            "State was treated as a multidimensional operating profile, not a single "
            "scalar. Axes included lapse-like error rate, drift, variability and "
            "reliability. Capacity was treated as a multidimensional resource profile, "
            "including selected hidden size, selection confidence, complexity preference, "
            "high-capacity advantage, load robustness and consistency.",
        ),
        (
            "Dynamics",
            "Hidden trajectories from fitted GRUs were summarized by participation "
            "ratio, covariance rank, trajectory radius, step norm, hidden variability, "
            "fixed-point/Jacobian summaries and latent decoder performance. These "
            "quantities were interpreted as fitted-model dynamics, not direct neural measurements.",
        ),
        (
            "Neurophysiology and external validation",
            "ds007554 EEG, fNIRS and ECG features were extracted and modeled against "
            "task load and repaired state/capacity coordinates. COG-BCI EEG features "
            "were extracted from EEGLAB marker trials. TU Berlin EEG and NIRS block "
            "windows were extracted from MATLAB archives and aligned to N-back load. "
            "HBN was used as scalability and developmental-exploratory evidence.",
        ),
        (
            "Statistical testing",
            "Analyses included leave-one-participant-out prediction, Spearman "
            "associations, OLS models with task/session/subject controls where "
            "appropriate, permutation tests, bootstrap confidence intervals and "
            "Benjamini-Hochberg correction. Claim strength was graded as strong, "
            "moderate, exploratory, negative or failed validation based on convergence, "
            "correction, controls and preregistered gate outcomes.",
        ),
    ]
    for title, text in methods:
        add_heading(doc, title, 2)
        add_para(doc, text)

    add_heading(doc, "Limitations", 1)
    limitations = [
        "The ANN gate failed for residualized state recovery, so state claims must remain exploratory.",
        "ds007554 correctness was reconstructed from push-button timing, not provided as explicit trial-level correctness labels.",
        "The ds007554 behavioral descriptive baseline outperformed the coordinate models.",
        "Capacity is participant-level in the current implementation, limiting formal session-level capacity test-retest claims.",
        "HBN is useful for scalability but not direct adult N-back/PVT evidence; RT is absent in the unified table and correctness is task-limited.",
        "Recurrent geometry is fitted-model geometry, not direct neural geometry.",
        "Vanilla RNN and LSTM architecture variants were not trained in the current full run, so architecture-variant robustness cannot be claimed.",
        "Neurophysiology analyses use derived features rather than full encoding models or source-localized neural measures.",
    ]
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Data and code availability", 1)
    add_para(
        doc,
        "All tables, figures, audit outputs and manuscript source fragments used in "
        "this draft are stored under state_capacity_tinyrnn/outputs. The Word "
        "manuscript was generated by state_capacity_tinyrnn/scripts/"
        "build_nature_style_word_manuscript.py.",
    )


def add_figures(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Figures", 1)
    for filename, caption in FIGURE_CAPTIONS.items():
        fig_path = FIGS / filename
        if not fig_path.exists():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_path), width=Inches(6.2))
        add_caption(doc, caption)

    add_heading(doc, "Figure Source Data Map", 1)
    rows = [
        ["Figure 1", "event_counts_by_dataset.csv; master_effects_summary.csv; claim_audit.tsv"],
        ["Figure 2", "ann_intervention_gate_results.csv; ann_hybrid_recovery.csv; performance_matched_agent_pairs.csv"],
        ["Figure 3", "ds007554_discovery_model_comparison.csv; ds007554_bootstrap_effects.csv"],
        ["Figure 4", "recurrent_dynamics_state_capacity_tests.csv; latent_decoder_results.csv"],
        ["Figure 5", "ds007554_neurophys_models.csv; ds007554_neurophys_features.csv"],
        ["Figure 6", "cog_bci_validation_models.csv; cog_bci_eeg_features.csv"],
        ["Figure 7", "tu_berlin_load_validation.csv; hbn_scalability_tests.csv"],
        ["Figure 8", "robustness_master_table.csv; falsification_tests.csv"],
    ]
    add_table(
        doc,
        "Table 6 | Figure source-data index.",
        "outputs/tables and outputs/audit",
        ["Figure", "Primary source data files"],
        rows,
        [1300, 8060],
    )


def build() -> Path:
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_title_block(doc)
    add_main_text(doc)
    add_results(doc)
    add_discussion(doc)
    add_methods(doc)
    add_figures(doc)
    doc.save(DOCX_OUT)
    return DOCX_OUT


if __name__ == "__main__":
    path = build()
    print(path)
