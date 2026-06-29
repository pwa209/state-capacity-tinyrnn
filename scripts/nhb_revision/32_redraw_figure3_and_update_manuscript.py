from __future__ import annotations

import csv
import datetime as dt
import shutil
import sys
import zipfile
from pathlib import Path

BUNDLED_SITE = Path(
    "C:/Users/Gebruiker/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/Lib/site-packages"
)
DOCUMENT_SKILL_SCRIPTS = Path(
    "C:/Users/Gebruiker/.codex/plugins/cache/openai-primary-runtime/documents/26.601.10930/skills/documents/scripts"
)
if BUNDLED_SITE.exists() and str(BUNDLED_SITE) not in sys.path:
    sys.path.append(str(BUNDLED_SITE))
if DOCUMENT_SKILL_SCRIPTS.exists() and str(DOCUMENT_SKILL_SCRIPTS) not in sys.path:
    sys.path.append(str(DOCUMENT_SKILL_SCRIPTS))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
TABLE_DIR = ROOT / "outputs" / "nhb_revision" / "tables"
OUT_DIR = ROOT / "outputs" / "nhb_revision" / "display_item_revision"
OUT_DIR.mkdir(parents=True, exist_ok=True)

MANUSCRIPT_IN = Path("C:/Users/Gebruiker/Downloads/state_capacity_manuscript_revised_9.docx")
CLEAN_DOCX = OUT_DIR / "state_capacity_manuscript_revised_9_clean_updated_fig3.docx"
TRACKED_DOCX = OUT_DIR / "state_capacity_manuscript_revised_9_tracked_changes_updated_fig3.docx"

FIG_BASE = OUT_DIR / "fig3_artificial_agent_gates_revised"
FIG_PNG = FIG_BASE.with_suffix(".png")
FIG_PDF = FIG_BASE.with_suffix(".pdf")
FIG_SVG = FIG_BASE.with_suffix(".svg")
FIG_SOURCE_XLSX = OUT_DIR / "fig3_revised_source_data.xlsx"
CHANGE_LOG = OUT_DIR / "figure3_manuscript_revision_log.md"

MODEL_ORDER = ["vanilla_rnn", "gru", "lstm"]
MODEL_LABEL = {"vanilla_rnn": "RNN", "gru": "GRU", "lstm": "LSTM"}
AXIS_LABEL = {"state_severity": "State", "capacity_level": "Capacity"}
FEATURE_LABEL = {"raw_fingerprint": "Raw", "residualized_fingerprint": "Residualized"}


def p_text(p: float) -> str:
    if p < 0.001:
        return "P<0.001"
    return f"P={p:.3f}"


def load_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    family = pd.read_csv(TABLE_DIR / "architecture_perturbation_gate_results.csv")
    hybrid = pd.read_csv(TABLE_DIR / "architecture_hybrid_recovery_results.csv")
    loo = pd.read_csv(TABLE_DIR / "leave_one_architecture_gate_results.csv")
    return family, hybrid, loo


def draw_revised_figure3() -> None:
    family, hybrid, loo = load_tables()

    family_resid = (
        family.query("feature_set == 'residualized_fingerprint' and task_family == 'overall'")
        .set_index("model_family")
        .loc[MODEL_ORDER]
        .reset_index()
    )
    loo_resid = (
        loo.query("feature_set == 'residualized_fingerprint' and task_family == 'overall'")
        .set_index("heldout_model_family")
        .loc[MODEL_ORDER]
        .reset_index()
    )
    hybrid = hybrid.copy()
    hybrid["model_family"] = pd.Categorical(hybrid["model_family"], MODEL_ORDER, ordered=True)
    hybrid["axis"] = hybrid["target_axis"].map(AXIS_LABEL)
    hybrid["feature_label"] = hybrid["feature_set"].map(FEATURE_LABEL)
    hybrid = hybrid.sort_values(["feature_set", "model_family", "target_axis"])

    with pd.ExcelWriter(FIG_SOURCE_XLSX, engine="openpyxl") as writer:
        family_resid.to_excel(writer, sheet_name="panel_a_family_gate", index=False)
        loo_resid.to_excel(writer, sheet_name="panel_b_leave_one_arch", index=False)
        hybrid.to_excel(writer, sheet_name="panel_c_d_scalar_recovery", index=False)

    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 8,
            "axes.linewidth": 0.7,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "xtick.major.width": 0.7,
            "ytick.major.width": 0.7,
            "pdf.fonttype": 42,
            "svg.fonttype": "none",
        }
    )

    colors = {
        "State": "#0072B2",
        "Capacity": "#D55E00",
        "bar": "#4D4D4D",
        "null": "#B7B7B7",
        "pass": "#2C7A4B",
        "fail": "#E8E1D9",
    }

    fig = plt.figure(figsize=(12.4, 8.2), dpi=300)
    gs = fig.add_gridspec(2, 2, width_ratios=[1.08, 1.28], height_ratios=[0.9, 1.15], wspace=0.46, hspace=0.50)

    ax_a = fig.add_subplot(gs[0, 0])
    x = np.arange(len(MODEL_ORDER))
    ba = family_resid["balanced_accuracy"].to_numpy()
    null95 = family_resid["null_95th_percentile"].to_numpy()
    ax_a.bar(x, ba, color=colors["bar"], width=0.58)
    ax_a.scatter(x, null95, color=colors["null"], edgecolor="black", linewidth=0.4, marker="D", s=28, zorder=3)
    ax_a.axhline(0.5, color="#999999", linewidth=0.8, linestyle=(0, (3, 2)))
    for i, row in family_resid.iterrows():
        ax_a.text(i, row["balanced_accuracy"] + 0.018, f"{row['balanced_accuracy']:.2f}", ha="center", va="bottom", fontsize=8)
        ax_a.text(i, row["null_95th_percentile"] - 0.042, "null\n95%", ha="center", va="top", fontsize=6.8, color="#222222")
    ax_a.set_xticks(x)
    ax_a.set_xticklabels([MODEL_LABEL[m] for m in MODEL_ORDER])
    ax_a.set_ylim(0.45, 1.02)
    ax_a.set_ylabel("Balanced accuracy")
    ax_a.set_title("a  Residualized intervention-family gate", loc="left", fontweight="bold", fontsize=9)

    ax_b = fig.add_subplot(gs[0, 1])
    x = np.arange(len(MODEL_ORDER))
    ax_b.bar(x, loo_resid["balanced_accuracy"], color="#6B6B6B", width=0.56)
    ax_b.scatter(x, loo_resid["auc"], color="white", edgecolor="black", linewidth=0.7, marker="o", s=34, zorder=3, label="AUC")
    ax_b.axhline(0.5, color="#999999", linewidth=0.8, linestyle=(0, (3, 2)))
    for i, row in loo_resid.iterrows():
        ax_b.text(i, row["balanced_accuracy"] - 0.028, f"BA {row['balanced_accuracy']:.2f}", ha="center", va="top", fontsize=8, color="white")
        ax_b.text(i, row["auc"] + 0.018, f"AUC {row['auc']:.2f}", ha="center", va="bottom", fontsize=7.5)
    ax_b.set_xticks(x)
    ax_b.set_xticklabels([MODEL_LABEL[m] for m in MODEL_ORDER])
    ax_b.set_ylim(0.45, 1.02)
    ax_b.set_ylabel("Held-out performance")
    ax_b.set_title("b  Leave-one-architecture generalization", loc="left", fontweight="bold", fontsize=9)

    ax_c = fig.add_subplot(gs[1, 0])
    xticks = []
    xticklabels = []
    xpos = {}
    gap = 0.65
    base = 0.0
    for f_i, feature in enumerate(["raw_fingerprint", "residualized_fingerprint"]):
        for m_i, model in enumerate(MODEL_ORDER):
            pos = base + m_i
            xpos[(feature, model)] = pos
            xticks.append(pos)
            xticklabels.append(MODEL_LABEL[model])
        base += len(MODEL_ORDER) + gap

    for axis, jitter in [("State", -0.08), ("Capacity", 0.08)]:
        rows = hybrid[hybrid["axis"] == axis]
        for _, row in rows.iterrows():
            pos = xpos[(row["feature_set"], row["model_family"])] + jitter
            marker = "o" if row["pass_gate"] else "x"
            ax_c.scatter(
                pos,
                row["spearman_rho"],
                s=58,
                color=colors[axis],
                marker=marker,
                linewidth=1.5,
                zorder=4,
                label=axis if row["feature_set"] == "raw_fingerprint" and row["model_family"] == "gru" else None,
            )
            if (
                row["model_family"] == "gru"
                and row["target_axis"] == "capacity_level"
                and row["feature_set"] == "residualized_fingerprint"
            ):
                ax_c.annotate(
                    "GRU capacity:\nraw pass,\nresidualized fail",
                    (pos, row["spearman_rho"]),
                    xytext=(16, 18),
                    textcoords="offset points",
                    fontsize=7,
                    arrowprops=dict(arrowstyle="-", color="#555555", linewidth=0.6),
                    ha="left",
                    va="bottom",
                )
    ax_c.axhline(0.4, color="#333333", linewidth=0.8, linestyle=(0, (3, 2)))
    ax_c.axhline(-0.4, color="#333333", linewidth=0.8, linestyle=(0, (3, 2)))
    ax_c.axhline(0, color="#BBBBBB", linewidth=0.7)
    ax_c.text(np.mean([xpos[("raw_fingerprint", m)] for m in MODEL_ORDER]), -0.63, "Raw fingerprints", ha="center", va="top", fontsize=8)
    ax_c.text(np.mean([xpos[("residualized_fingerprint", m)] for m in MODEL_ORDER]), -0.63, "Residualized fingerprints", ha="center", va="top", fontsize=8)
    ax_c.set_xticks(xticks)
    ax_c.set_xticklabels(xticklabels)
    ax_c.set_ylim(-0.7, 1.02)
    ax_c.set_ylabel("Spearman rho with imposed magnitude")
    ax_c.set_title("c  Hybrid scalar recovery: raw versus residualized", loc="left", fontweight="bold", fontsize=9)
    ax_c.legend(frameon=False, loc="upper left", fontsize=8, handletextpad=0.3)
    ax_c.text(xpos[("raw_fingerprint", "vanilla_rnn")] - 0.45, 0.43, "gate", fontsize=7, color="#333333")
    ax_c.text(xpos[("raw_fingerprint", "vanilla_rnn")] - 0.45, -0.37, "gate", fontsize=7, color="#333333")

    ax_d = fig.add_subplot(gs[1, 1])
    matrix_rows = []
    for model in MODEL_ORDER:
        for target in ["state_severity", "capacity_level"]:
            matrix_rows.append((model, target))
    matrix = np.zeros((len(matrix_rows), 2))
    labels = []
    row_labels = []
    for r, (model, target) in enumerate(matrix_rows):
        row_labels.append(f"{MODEL_LABEL[model]} {AXIS_LABEL[target]}")
        row_labels[-1] = row_labels[-1].replace("Capacity", "Cap.")
        label_row = []
        for c, feature in enumerate(["raw_fingerprint", "residualized_fingerprint"]):
            row = hybrid[
                (hybrid["model_family"] == model)
                & (hybrid["target_axis"] == target)
                & (hybrid["feature_set"] == feature)
            ].iloc[0]
            passed = bool(row["pass_gate"])
            matrix[r, c] = 1 if passed else 0
            label_row.append(f"{row['spearman_rho']:.2f}\n{p_text(float(row['nominal_p_value']))}\n{'PASS' if passed else 'FAIL'}")
        labels.append(label_row)

    from matplotlib.colors import ListedColormap

    ax_d.imshow(matrix, cmap=ListedColormap([colors["fail"], colors["pass"]]), vmin=0, vmax=1, aspect="auto")
    ax_d.set_xticks([0, 1])
    ax_d.set_xticklabels(["Raw", "Residualized"])
    ax_d.set_yticks(np.arange(len(row_labels)))
    ax_d.set_yticklabels(row_labels)
    ax_d.set_title("d  Scalar-coordinate gate matrix", loc="left", fontweight="bold", fontsize=9)
    for r in range(matrix.shape[0]):
        for c in range(matrix.shape[1]):
            color = "white" if matrix[r, c] == 1 else "#222222"
            ax_d.text(c, r, labels[r][c], ha="center", va="center", fontsize=7, color=color)
    ax_d.set_xticks(np.arange(-0.5, 2, 1), minor=True)
    ax_d.set_yticks(np.arange(-0.5, len(row_labels), 1), minor=True)
    ax_d.grid(which="minor", color="white", linewidth=1.5)
    ax_d.tick_params(which="minor", bottom=False, left=False)
    for spine in ax_d.spines.values():
        spine.set_visible(False)

    fig.text(
        0.01,
        0.005,
        "Figure 3 revised from architecture perturbation and hybrid recovery tables. Filled circles indicate gates passed; x marks failed gates. "
        "Residualized scalar-coordinate evidence requires |rho| > 0.40 and P < 0.05.",
        fontsize=7,
        color="#333333",
    )
    fig.savefig(FIG_PNG, bbox_inches="tight")
    fig.savefig(FIG_PDF, bbox_inches="tight")
    fig.savefig(FIG_SVG, bbox_inches="tight")
    plt.close(fig)


def paragraph_replacements(paragraphs: list[str]) -> dict[int, str]:
    repl: dict[int, str] = {}
    repl[2] = (
        "The momentary state of a cognitive system, such as its arousal or response policy, has long been distinguished from its enduring capacity, "
        "such as the size of the workspace it can hold in mind. These two kinds are separated within individual paradigms, but whether they form one "
        "paradigm-general two-axis structure is untested, because human manipulations confound them and lack an external standard. We use recurrent "
        "neural networks, in which state-like interventions leave the computational graph intact while capacity-like interventions alter it, to build "
        "such a standard. Across vanilla RNN, GRU and LSTM agents the two intervention families are separable and architecture-robust. Projecting "
        "289,827 behavioural events from four open human datasets onto these axes, capacity validated as a convergent, load-sensitive human profile, "
        "whereas neither axis provided general residualized scalar-coordinate recovery in the hybrid-agent test. State and capacity are separable "
        "families of computation, but the present evidence supports capacity at the profile level and requires qualified language for scalar coordinates."
    )
    repl[10] = (
        "The results that follow support a deliberately bounded conclusion. The state and capacity intervention families separate cleanly and "
        "architecture-robustly in the networks, and capacity validates in humans as a convergent profile with geometric, state-controlled load-pressure "
        "and generalisation evidence. The corresponding state signal is visible in raw network fingerprints but does not generally survive residualisation "
        "against capacity and architecture; importantly, residualised scalar capacity recovery in the hybrid agents is also not robust across architectures. "
        "We therefore frame the contribution as a constructive falsification. The state and capacity distinction is real enough to separate artificial "
        "perturbation families, to support capacity as a transferable human profile, and to organise human behaviour beyond baseline, but the strong "
        "hypothesis of two architecture-free scalar coordinates is not supported by the present scalar-recovery gate."
    )
    repl[19] = "Intervention families separate cleanly, but residualised scalar recovery is not robust"
    repl[21] = (
        "Fig. 3 | Intervention families separate cleanly, but residualised scalar-coordinate recovery is not robust. "
        "a, Balanced accuracy of the intervention-family classifier on residualised fingerprints within each recurrent family, with permutation-null "
        "95th-percentile benchmarks. b, Leave-one-architecture-out generalisation for residualised fingerprints, showing that the family boundary "
        "transfers across recurrent architectures. c, Hybrid scalar recovery is strong for both axes on raw fingerprints, including capacity in GRU "
        "agents, but residualised recovery is architecture-limited for state and fails the pre-specified gate for capacity in all three recurrent families. "
        "d, Gate matrix for raw and residualised scalar recovery. The residualised scalar-coordinate criterion was |Spearman rho| > 0.40 with P < 0.05."
    )
    repl[23] = (
        "Whether the two families reduce to two scalar severities is a stronger question, and here the result is the pivot of the study. Hybrid agents "
        "carrying graded mixtures of state and capacity interventions let us test scalar recovery, that is, whether a single projected coordinate tracks "
        "the underlying intervention magnitude. In raw fingerprints, recovery was strong for both axes (median Spearman rho = 0.73 for state, 0.93 for "
        "capacity; Fig. 3c,d; Extended Data Table 1). But raw fingerprints conflate the targeted intervention with architecture and with the other axis. "
        "Once we residualised against both, which is the operation any claim of an architecture-free coordinate must withstand, scalar state recovery was "
        "not robust (median rho = 0.12, pass in one of three architectures) and scalar capacity recovery in hybrid agents also failed the residualised "
        "gate (median rho = -0.02, pass in zero of three architectures). Thus GRU capacity passed the raw recovery test but not the residualised scalar "
        "coordinate test. The separability of the families is robust, but the present hybrid-agent evidence does not justify treating either axis as a "
        "general architecture-free scalar coordinate. We report this as a falsification rather than bury it: the same residualisation that the family "
        "classifier passes, the scalar-coordinate gate does not generally pass, and this dissociation governs the cautious language we use throughout."
    )
    repl[38] = paragraphs[38].replace(
        "capacity validation, two moderate claims",
        "capacity profile validation, two moderate claims",
    ).replace(
        "The pre-specified failure criterion for a scalar state coordinate was met, while the criteria for family separability and capacity validation were passed.",
        "The pre-specified residualised scalar-coordinate gate was not generally met, while the criteria for family separability and capacity profile validation were passed.",
    )
    repl[41] = (
        "Using recurrent neural networks as an external standard, we asked whether the state and capacity distinction that is orthogonalised within "
        "individual paradigms reflects a single, paradigm-general two-axis structure of cognitive variation5,6. Three findings define the answer. "
        "The distinction is real at the level of intervention families, because graph-preserving and graph-altering perturbations leave separable, "
        "architecture-robust signatures and human behaviour is organised in part by these two families. Capacity transfers as a profile-level human "
        "construct, with convergent recurrent-geometry, working-memory load-pressure and generalisation evidence. But a clean scalar coordinate does "
        "not transfer for either axis under the residualised hybrid-agent gate; state, in particular, behaves in humans as a session-bound and task-bound "
        "reliability profile. The two-axis intuition that runs from Chomsky and Cronbach to modern process models1,2,7 is therefore vindicated as a "
        "structure of separable families and supported for capacity at the profile level, but its strongest form, two architecture-free scalar coordinates "
        "in humans, is rejected by the present evidence."
    )
    repl[42] = (
        "Why should capacity yield a profile-level validation while scalar recovery remains fragile? The asymmetry is informative rather than disappointing. "
        "Capacity, as a structural resource, leaves a comparatively stereotyped imprint: a larger workspace expands the effective dimensionality of the "
        "recurrent trajectory55,56,57 and confers a load-robustness that load itself makes visible16, and these consequences are similar across architectures. "
        "However, the hybrid-agent audit shows that this convergent profile should not be confused with a residualised scalar-coordinate pass. State, as an "
        "operating regime, is realised differently in different systems and is entangled with capacity in any uncontrolled sample, because the same lapse rate "
        "or drift can arise from changes in gain, noise or threshold34,35, and because humans compensate for momentary state with strategy in ways that "
        "networks under a fixed policy do not. When fingerprints are residualised to remove architecture and the complementary axis, the shared variance that "
        "made raw scalar recovery look strong is precisely what is removed. The practical implication is that state is better modelled as a multidimensional "
        "operating profile, spanning lapse, drift, variability and reliability, and capacity as a convergent resource profile, until a prospective dataset "
        "validates either as a stable scalar quantity. The trial-count dependence we observe, and its resonance with the reliability paradox39,40,41,42, points "
        "to the same prescription."
    )
    repl[45] = paragraphs[45].replace(
        "Capacity was estimated at the participant level and is stable between persons partly by construction, so its validation rests on convergent geometry, load pressure and held-out generalisation rather than on within-person variation.",
        "Capacity was estimated at the participant level and is stable between persons partly by construction, so its validation rests on convergent geometry, load pressure and held-out generalisation rather than on within-person variation or on a passed residualised scalar-coordinate gate.",
    ).replace(
        "Most importantly, the failure of residualised scalar state recovery is a falsification of a specific strong claim, not evidence that no state coordinate exists; it indicates that the present data are insufficient to validate one.",
        "Most importantly, the failure of residualised scalar recovery is a falsification of a specific strong claim, not evidence that no state or capacity coordinate exists; it indicates that the present data are insufficient to validate a general scalar coordinate."
    )
    repl[70] = (
        "Two operations protect the central comparisons from trivial confounds. First, agents were matched on mean accuracy across intervention families, "
        "with matching diagnostics retained, so that family membership could not be read off performance. Second, fingerprints were residualised by "
        "regressing out architecture family and the complementary axis, isolating the variance specific to the targeted intervention family. Raw fingerprints "
        "were retained for descriptive recovery, but the residualised fingerprint is the appropriate substrate for any claim of an architecture-free coordinate, "
        "because it removes precisely the shared variance that would otherwise inflate recovery. The central scalar-recovery result of the study is defined "
        "by this residualisation: raw recovery is strong, whereas residualised recovery is not generally robust for either state or capacity in the hybrid-agent test."
    )
    repl[74] = (
        "For hybrid agents, a projected scalar coordinate, state or capacity, was correlated with the known underlying intervention magnitude using the Spearman "
        "rank correlation with bootstrap confidence intervals, computed separately on raw and on residualised fingerprints. The pre-specified gate for accepting "
        "a construct as a scalar coordinate was significant residualised recovery with |Spearman rho| > 0.40 and P < 0.05. Raw-only recovery, which conflates the "
        "targeted intervention with architecture and with the other axis, was defined in advance to be insufficient evidence for an architecture-free scalar axis. "
        "The failure of residualised scalar recovery to pass generally, despite the success of the family separation that uses the same residualisation, is reported "
        "as a constructive falsification rather than as a null finding."
    )
    repl[90] = paragraphs[90].replace(
        "capacity geometry and physiology are described as convergent and bounded respectively",
        "capacity geometry and physiology are described as convergent and bounded respectively, and capacity is not described as a residualised hybrid scalar-coordinate pass",
    )
    repl[180] = repl[21]
    return repl


def insert_figure_after(paragraph, image_path: Path, width_inches: float = 6.45) -> None:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    image_para = Paragraph(new_p, paragraph._parent)
    image_para.alignment = 1
    run = image_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def replace_para_text(paragraph, text: str) -> None:
    paragraph.text = text


def write_clean_docx() -> None:
    doc = Document(str(MANUSCRIPT_IN))
    originals = [p.text for p in doc.paragraphs]
    repl = paragraph_replacements(originals)
    for idx, new_text in repl.items():
        replace_para_text(doc.paragraphs[idx], new_text)
    insert_figure_after(doc.paragraphs[21], FIG_PNG)
    doc.save(CLEAN_DOCX)


def write_tracked_docx() -> None:
    base = OUT_DIR / "_tracked_base_with_fig3.docx"
    doc = Document(str(MANUSCRIPT_IN))
    insert_figure_after(doc.paragraphs[21], FIG_PNG)
    doc.save(base)
    originals = [p.text for p in Document(str(MANUSCRIPT_IN)).paragraphs]
    replacements = paragraph_replacements(originals)
    patch_full_paragraph_redlines(base, TRACKED_DOCX, replacements)
    base.unlink(missing_ok=True)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def element_text(p: etree._Element) -> str:
    return "".join(p.xpath(".//w:t/text()", namespaces=NS))


def enable_track_revisions(settings_root: etree._Element) -> None:
    if settings_root.find("w:trackRevisions", namespaces=NS) is None:
        settings_root.insert(0, etree.Element(w("trackRevisions")))


def next_change_id(doc_root: etree._Element) -> int:
    ids: list[int] = []
    for el in doc_root.xpath(".//*[@w:id]", namespaces=NS):
        try:
            ids.append(int(el.get(w("id"))))
        except Exception:
            pass
    return max(ids, default=0) + 1


def revision_run(tag: str, text_tag: str, text: str, cid: int, when: str) -> etree._Element:
    el = etree.Element(w(tag))
    el.set(w("id"), str(cid))
    el.set(w("author"), "Codex")
    el.set(w("date"), when)
    r = etree.SubElement(el, w("r"))
    t = etree.SubElement(r, w(text_tag))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return el


def redline_paragraph(p: etree._Element, old_text: str, new_text: str, cid: int, when: str) -> int:
    ppr = p.find("w:pPr", namespaces=NS)
    for child in list(p):
        if child is not ppr:
            p.remove(child)
    p.append(revision_run("del", "delText", old_text, cid, when))
    p.append(revision_run("ins", "t", new_text, cid + 1, when))
    return cid + 2


def patch_full_paragraph_redlines(base_docx: Path, out_docx: Path, replacements: dict[int, str]) -> None:
    when = dt.datetime.now(dt.UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    with zipfile.ZipFile(base_docx, "r") as zin:
        overrides: dict[str, bytes] = {}
        doc_root = etree.fromstring(zin.read("word/document.xml"))
        settings_name = "word/settings.xml"
        settings_root = (
            etree.fromstring(zin.read(settings_name))
            if settings_name in zin.namelist()
            else etree.Element(w("settings"))
        )
        enable_track_revisions(settings_root)

        body_paras = doc_root.xpath("./w:body/w:p", namespaces=NS)
        cid = next_change_id(doc_root)
        changed = 0
        for idx, new_text in replacements.items():
            body_idx = idx + 1 if idx > 21 else idx
            if body_idx >= len(body_paras):
                continue
            p = body_paras[body_idx]
            old_text = element_text(p)
            if not old_text:
                continue
            cid = redline_paragraph(p, old_text, new_text, cid, when)
            changed += 1

        overrides["word/document.xml"] = etree.tostring(
            doc_root, xml_declaration=True, encoding="UTF-8", standalone="yes"
        )
        overrides[settings_name] = etree.tostring(
            settings_root, xml_declaration=True, encoding="UTF-8", standalone="yes"
        )

        with zipfile.ZipFile(out_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                name = info.filename
                if name in overrides:
                    zout.writestr(name, overrides[name])
                else:
                    zout.writestr(info, zin.read(name))
    print(f"[OK] wrote full redline {out_docx} (paragraphs={changed})")


def write_change_log() -> None:
    _, hybrid, _ = load_tables()
    gru_cap = hybrid[(hybrid["model_family"] == "gru") & (hybrid["target_axis"] == "capacity_level")]
    raw = gru_cap[gru_cap["feature_set"] == "raw_fingerprint"].iloc[0]
    resid = gru_cap[gru_cap["feature_set"] == "residualized_fingerprint"].iloc[0]
    text = f"""# Figure 3 and manuscript correction log

## Core audit result

- GRU capacity raw hybrid recovery: Spearman rho = {raw['spearman_rho']:.3f}, P = {raw['nominal_p_value']:.3g}, gate = {raw['pass_gate']}.
- GRU capacity residualized hybrid recovery: Spearman rho = {resid['spearman_rho']:.3f}, P = {resid['nominal_p_value']:.3g}, gate = {resid['pass_gate']}.
- Therefore, the statement "capacity GRU passed" is true for the raw hybrid recovery gate only, not for the residualized scalar-coordinate gate.

## Manuscript adjustment

The revised manuscript now distinguishes:

1. strong residualized intervention-family separability;
2. strong raw hybrid scalar recovery for both axes;
3. failed general residualized scalar-coordinate recovery for capacity and state in the hybrid-agent test;
4. capacity support as a convergent, profile-level human construct rather than as a passed residualized scalar coordinate.

## Outputs

- Independent Figure 3: `{FIG_PNG.name}`, `{FIG_PDF.name}`, `{FIG_SVG.name}`
- Figure 3 source data workbook: `{FIG_SOURCE_XLSX.name}`
- Clean manuscript with updated Figure 3: `{CLEAN_DOCX.name}`
- Tracked-changes manuscript with updated Figure 3: `{TRACKED_DOCX.name}`
"""
    CHANGE_LOG.write_text(text, encoding="utf-8")


def main() -> None:
    if not MANUSCRIPT_IN.exists():
        raise FileNotFoundError(MANUSCRIPT_IN)
    draw_revised_figure3()
    write_clean_docx()
    write_tracked_docx()
    write_change_log()
    print(f"Wrote {FIG_PNG}")
    print(f"Wrote {FIG_PDF}")
    print(f"Wrote {FIG_SVG}")
    print(f"Wrote {FIG_SOURCE_XLSX}")
    print(f"Wrote {CLEAN_DOCX}")
    print(f"Wrote {TRACKED_DOCX}")
    print(f"Wrote {CHANGE_LOG}")


if __name__ == "__main__":
    main()
