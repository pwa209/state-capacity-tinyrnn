from __future__ import annotations

import json
import re
import sys
import textwrap
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parents[1]
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from nhb_utils import NHB, NHB_AUDIT, NHB_MANUSCRIPT, NHB_TABLES, TABLES, append_manifest, append_registry, ensure_nhb_dirs


ANALYSIS_ID = "nhb_26_record_reference_package"
SCRIPT_NAME = "scripts/nhb_revision/26_build_record_reference_package.py"
RECORD_DIR = NHB / "record_package"
REPORT_DOCX = RECORD_DIR / "state_capacity_full_results_record_report.docx"
REPORT_MD = RECORD_DIR / "state_capacity_full_results_record_report.md"
EXCEL_OUT = RECORD_DIR / "state_capacity_all_results_workbook.xlsx"
PY_REF = RECORD_DIR / "state_capacity_method_analysis_reference.py"
MANIFEST_JSON = RECORD_DIR / "state_capacity_record_manifest.json"
MAX_EXCEL_ROWS = 1_048_000
MAX_DOC_TABLE_ROWS = 20


PRESET = {
    "font": "Calibri",
    "ink": "14213D",
    "muted": "5C6672",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "header_fill": "F4F6F9",
    "body_size": 10.5,
    "caption_size": 8.5,
    "cell_margins": {"top": 80, "bottom": 80, "start": 120, "end": 120},
    "table_indent_dxa": 120,
}


KEY_TABLES = [
    "architecture_perturbation_gate_results.csv",
    "architecture_hybrid_recovery_results.csv",
    "leave_one_architecture_gate_results.csv",
    "state_split_half_reliability.csv",
    "state_bootstrap_reliability.csv",
    "state_capacity_variance_decomposition.csv",
    "capacity_pressure_models.csv",
    "capacity_variant_validation_summary.csv",
    "incremental_value_model_comparison.csv",
    "physiology_robustness_models.csv",
    "physiology_permutation_controls.csv",
]


STEP_PROTOCOL = [
    ("00", "Freeze current state", "Hash and inventory current scripts, outputs and data before revision."),
    ("01", "Architecture variants", "Train/check vanilla RNN, GRU and LSTM human architecture variants where feasible."),
    ("02", "True architecture artificial agents", "Train artificial vanilla RNN, GRU and LSTM agents and test state/capacity perturbation gates."),
    ("03", "State reliability", "Split-half and bootstrap reliability of state profiles by dataset, task and trial-count bin."),
    ("04", "State early-late prediction", "Use early-session state/recent behaviour to predict later outcomes."),
    ("05", "Variance decomposition", "Estimate whether state behaves as session/task profile and capacity as participant-level profile."),
    ("06", "Capacity pressure", "Model TU Berlin load, capacity and load-by-capacity effects with state covariates."),
    ("07", "Capacity transfer", "Leave-one-task capacity prediction and shuffled capacity controls."),
    ("08", "Capacity ablation", "Compare full, model-only, behaviour-only and ablated capacity definitions."),
    ("09", "Incremental value", "Compare task/dataset, behavioural baseline, state/capacity and shuffled/random controls."),
    ("10", "Low-data curves", "Assess calibration/data sufficiency for profile estimation."),
    ("11", "Leave-one-dataset validation", "Check direction consistency when each dataset is held out."),
    ("12", "Leave-one-task validation", "Check direction consistency when each task family is held out."),
    ("13", "Feature taxonomy", "Freeze state/capacity/control feature taxonomy and loading enrichment."),
    ("14", "Physiology controls", "Run EEG/fNIRS/ECG row-level association screens and blocked permutation controls."),
    ("15", "Interaction decomposition", "Separate general state-capacity interaction from load-by-capacity pressure."),
    ("16", "Leakage audit", "Record split rules and leakage checks for state/capacity predictors."),
    ("17", "Shuffle controls", "Report random-axis, shuffled-coordinate and profile shuffle falsification tests."),
    ("18", "Display items", "Regenerate Nature-style figures and source-data copies."),
    ("19", "Fancy figures", "Generate manuscript-facing NHB/NMI-style figures and source data."),
    ("20", "Manuscript", "Build revised NHB/NMI manuscript from proposal intro/methods and updated results."),
    ("21", "Record package", "Generate full results report, Excel workbook and Python method-analysis reference."),
]


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.strip("#")
    return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:], 16))


def read_table(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    sep = "\t" if path.suffix.lower() == ".tsv" else ","
    try:
        return pd.read_csv(path, sep=sep)
    except Exception:
        return pd.DataFrame()


def fmt_num(value: object, digits: int = 3) -> str:
    try:
        x = float(value)
    except Exception:
        return "NA"
    if not np.isfinite(x):
        return "NA"
    if abs(x) < 0.001 and x != 0:
        return f"{x:.2e}"
    if abs(x) < 1:
        return f"{x:.{digits}g}"
    return f"{x:.{digits}f}".rstrip("0").rstrip(".")


def fmt_p(value: object) -> str:
    try:
        x = float(value)
    except Exception:
        return "NA"
    if not np.isfinite(x):
        return "NA"
    if x < 0.001:
        return f"{x:.2e}"
    return f"{x:.3f}".rstrip("0").rstrip(".")


def rel(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except Exception:
        return path.as_posix()


def collect_result_files() -> pd.DataFrame:
    roots = [
        ("base_table", TABLES),
        ("nhb_table", NHB_TABLES),
        ("nhb_audit", NHB_AUDIT),
        ("fancy_source", NHB / "fancy_figures" / "source_data"),
        ("manuscript_source", NHB_MANUSCRIPT),
    ]
    rows = []
    for kind, root in roots:
        if not root.exists():
            continue
        for path in sorted(root.glob("*")):
            if path.is_file() and path.suffix.lower() in {".csv", ".tsv", ".json", ".md", ".txt"}:
                n_rows = ""
                n_cols = ""
                if path.suffix.lower() in {".csv", ".tsv"}:
                    df = read_table(path)
                    n_rows = len(df)
                    n_cols = len(df.columns)
                rows.append(
                    {
                        "kind": kind,
                        "file": path.name,
                        "relative_path": rel(path),
                        "bytes": path.stat().st_size,
                        "last_write_time": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
                        "rows": n_rows,
                        "columns": n_cols,
                    }
                )
    return pd.DataFrame(rows)


def key_stats() -> dict[str, object]:
    counts = read_table(TABLES / "event_counts_by_dataset.csv")
    arch = read_table(NHB_TABLES / "architecture_perturbation_gate_results.csv")
    hybrid = read_table(NHB_TABLES / "architecture_hybrid_recovery_results.csv")
    leave = read_table(NHB_TABLES / "leave_one_architecture_gate_results.csv")
    state_var = read_table(NHB_TABLES / "state_capacity_variance_decomposition.csv")
    cap = read_table(NHB_TABLES / "capacity_pressure_models.csv")
    phys = read_table(NHB_TABLES / "physiology_robustness_models.csv")
    claims = read_table(NHB_AUDIT / "nhb_final_claim_audit.tsv")
    inc = read_table(NHB_TABLES / "incremental_value_model_comparison.csv")
    stats: dict[str, object] = {
        "total_events": int(counts["n_included_events"].sum()) if not counts.empty else np.nan,
        "total_subjects_naive_sum": int(counts["n_subjects"].sum()) if not counts.empty else np.nan,
        "datasets": "; ".join(counts["dataset"].astype(str)) if not counts.empty else "",
        "claims_strong": int((claims.get("claim_strength", pd.Series(dtype=str)) == "strong").sum()) if not claims.empty else 0,
        "claims_moderate": int((claims.get("claim_strength", pd.Series(dtype=str)) == "moderate").sum()) if not claims.empty else 0,
        "claims_qualified": int((claims.get("claim_strength", pd.Series(dtype=str)) == "qualified").sum()) if not claims.empty else 0,
        "claims_exploratory": int((claims.get("claim_strength", pd.Series(dtype=str)) == "exploratory").sum()) if not claims.empty else 0,
    }
    if not arch.empty:
        g = arch[(arch["task_family"] == "overall") & (arch["feature_set"] == "residualized_fingerprint")]
        stats["arch_min_ba"] = g["balanced_accuracy"].astype(float).min()
        stats["arch_max_ba"] = g["balanced_accuracy"].astype(float).max()
        stats["arch_max_perm_p"] = g["permutation_p"].astype(float).max()
    if not hybrid.empty:
        stats["hybrid_raw_state_median_rho"] = hybrid[(hybrid["analysis_type"] == "hybrid_state_axis") & (hybrid["feature_set"] == "raw_fingerprint")]["spearman_rho"].astype(float).median()
        stats["hybrid_resid_state_median_rho"] = hybrid[(hybrid["analysis_type"] == "hybrid_state_axis") & (hybrid["feature_set"] == "residualized_fingerprint")]["spearman_rho"].astype(float).median()
    if not leave.empty:
        l = leave[leave["feature_set"] == "residualized_fingerprint"]
        stats["leave_one_arch_min_ba"] = l["balanced_accuracy"].astype(float).min()
    if not state_var.empty:
        stats["state_median_icc"] = state_var[state_var["construct"] == "state"]["icc"].astype(float).median()
        stats["capacity_median_icc"] = state_var[state_var["construct"] == "capacity"]["icc"].astype(float).median()
    if not cap.empty:
        hit = cap[(cap["outcome"] == "mean_accuracy") & (cap["predictor"] == "load_x_capacity")]
        if not hit.empty:
            stats["tu_accuracy_load_x_capacity_beta"] = hit.iloc[0].get("estimate")
            stats["tu_accuracy_load_x_capacity_q"] = hit.iloc[0].get("q_value")
        hit = cap[(cap["outcome"] == "rt_median") & (cap["predictor"] == "load_x_capacity")]
        if not hit.empty:
            stats["tu_rt_load_x_capacity_beta"] = hit.iloc[0].get("estimate")
            stats["tu_rt_load_x_capacity_q"] = hit.iloc[0].get("q_value")
    if not phys.empty:
        stats["physiology_tests"] = len(phys)
        stats["physiology_perm_fdr_rows"] = int((pd.to_numeric(phys.get("permutation_q_value", pd.Series(dtype=float)), errors="coerce") < 0.05).sum())
    if not inc.empty:
        for name in ["task_dataset", "behavioral_descriptive", "additive_state_capacity", "state_capacity_interaction", "random_axis_control", "shuffled_coordinate_control"]:
            rows = inc[inc["model_name"].eq(name)].copy()
            preferred = rows[rows.get("dataset", "").eq("all_supervised_datasets_context")] if "dataset" in rows else pd.DataFrame()
            row = preferred.iloc[0] if not preferred.empty else (rows.iloc[0] if not rows.empty else None)
            if row is not None:
                stats[f"rmse_{name}"] = float(row["RMSE"])
    return stats


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for margin in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(sec, margin, Inches(1))
    normal = doc.styles["Normal"]
    normal.font.name = PRESET["font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal.font.color.rgb = rgb(PRESET["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color in [("Heading 1", 16, PRESET["blue"]), ("Heading 2", 13, PRESET["blue"]), ("Heading 3", 11.5, PRESET["dark_blue"])]:
        st = doc.styles[name]
        st.font.name = PRESET["font"]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
    if "Caption" in doc.styles:
        cap = doc.styles["Caption"]
        cap.font.name = PRESET["font"]
        cap.font.size = Pt(PRESET["caption_size"])
        cap.font.color.rgb = rgb(PRESET["muted"])
    header = sec.header.paragraphs[0]
    header.text = "State-capacity full results record"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def set_cell_text(cell, text: object, bold: bool = False, size: float = 8.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("" if pd.isna(text) else str(text))
    run.bold = bold
    run.font.name = PRESET["font"]
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_shading(cell, fill: str = "F4F6F9") -> None:
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(tbl, widths: Sequence[int]) -> None:
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = tbl._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), str(PRESET["table_indent_dxa"]))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            pr = cell._tc.get_or_add_tcPr()
            tc_w = pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            mar = pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                pr.append(mar)
            for side, val in PRESET["cell_margins"].items():
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), str(val))
                node.set(qn("w:type"), "dxa")


def mark_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_doc_table(doc: Document, title: str, df: pd.DataFrame, columns: Sequence[str], widths: Sequence[int], max_rows: int = MAX_DOC_TABLE_ROWS) -> None:
    doc.add_paragraph(title, style="Caption")
    show = df.loc[:, [c for c in columns if c in df.columns]].head(max_rows).copy()
    if show.empty:
        add_para(doc, "No rows available.")
        return
    tbl = doc.add_table(rows=len(show) + 1, cols=len(show.columns))
    set_table_geometry(tbl, widths[: len(show.columns)])
    mark_header(tbl.rows[0])
    for i, col in enumerate(show.columns):
        set_cell_text(tbl.rows[0].cells[i], col, bold=True, size=8.1)
        set_shading(tbl.rows[0].cells[i])
    for r_idx, (_, row) in enumerate(show.iterrows(), start=1):
        for c_idx, col in enumerate(show.columns):
            value = row[col]
            if isinstance(value, (float, np.floating)):
                value = fmt_num(value)
            set_cell_text(tbl.rows[r_idx].cells[c_idx], value, size=7.6)
    doc.add_paragraph("")


def narrative_sections(stats: dict[str, object]) -> list[tuple[str, list[str]]]:
    return [
        (
            "Executive Summary",
            [
                f"The current full analysis contains {stats.get('total_events', 'NA'):,} included behavioural events across {stats.get('datasets', 'the available datasets')}. The strongest results are true architecture-specific artificial-agent family separation and human capacity-pressure validation. State is meaningful as a reliability profile, but residualized scalar state recovery remains qualified.",
                f"Architecture separation is strong: residualized balanced accuracy across true recurrent families ranges from {fmt_num(stats.get('arch_min_ba'))} to {fmt_num(stats.get('arch_max_ba'))}, with maximum permutation p = {fmt_p(stats.get('arch_max_perm_p'))}. Leave-one-architecture generalization remains above chance.",
                f"Capacity pressure is strong in TU Berlin N-back: load-by-capacity predicts accuracy with beta = {fmt_num(stats.get('tu_accuracy_load_x_capacity_beta'))}, q = {fmt_p(stats.get('tu_accuracy_load_x_capacity_q'))}, and median RT with beta = {fmt_num(stats.get('tu_rt_load_x_capacity_beta'))}, q = {fmt_p(stats.get('tu_rt_load_x_capacity_q'))}.",
                f"Physiology is present but bounded: {stats.get('physiology_tests', 0)} EEG/fNIRS/ECG association tests were run, with {stats.get('physiology_perm_fdr_rows', 0)} rows surviving blocked permutation FDR. These are exploratory physiological alignments, not direct neural-coordinate evidence.",
            ],
        ),
        (
            "What This Record Package Contains",
            [
                "This record package is designed as a working archive rather than a submission manuscript. The Word report summarizes every analysis family, gives the core result interpretation, and lists representative rows from the main tables. The Excel workbook is the exhaustive row-level archive: it contains the NHB revision tables, base result tables, audit files and fancy-figure source data as separate sheets. The Python reference file is a single-file method and analysis map with helper functions for reloading result tables.",
                "The package therefore separates three use cases. For reading and explanation, use the Word report. For checking every row, use the Excel workbook. For reconstructing the analysis logic or giving another model/program a compact technical handoff, use the Python reference file.",
            ],
        ),
        (
            "Interpretive Bottom Line",
            [
                "The project now supports the language of machine-defined state and capacity profiles. It does not yet support the strongest language of fully validated scalar state and capacity coordinates in humans.",
                "Capacity is the cleaner construct: it shows architecture robustness, participant-level stability, geometry/load-pressure convergence and moderate ablation robustness.",
                "State remains scientifically useful but should be described as a session/task reliability or operating-state profile. Its strongest empirical support comes from reliability, within-person variability and COG-BCI-style behavioural instability, while scalar recovery remains the central failed/qualified gate.",
                "The current manuscript should therefore lead with a constructive falsification frame: the state/capacity distinction is real enough to separate artificial perturbation families and support capacity-pressure validation, but the state construct must be written with qualified language until a prospective dataset validates it as a stable scalar coordinate.",
            ],
        ),
    ]


def detailed_result_sections(stats: dict[str, object]) -> list[tuple[str, list[str]]]:
    return [
        (
            "Dataset Scale and Eligibility",
            [
                f"The full event inventory contains {stats.get('total_events', 'NA'):,} included events. This includes COG-BCI, ds007554, TU Berlin EEG-NIRS and HBN Release 4. The datasets differ sharply in what they can validate: COG-BCI is best for repeated-session behavioural state; TU Berlin is best for N-back load pressure; ds007554 is multimodal but requires reconstructed correctness; HBN provides large-scale scalability rather than close task replication.",
                "The main confirmatory path therefore does not treat all datasets as interchangeable. Dataset role is part of the interpretation: capacity pressure is anchored in TU Berlin, state reliability is strongest in COG-BCI and pooled reliability analyses, physiology is multimodal but exploratory, and HBN is primarily a scale/generalization stress test.",
            ],
        ),
        (
            "Artificial-Agent and Architecture Gate",
            [
                f"The key artificial-agent result is no longer approximate replay. The revised Step 02 trained true architecture-specific artificial agents for vanilla RNN, GRU and LSTM families. After matching performance and residualising fingerprints, the state-versus-capacity family classifier remained strong across recurrent families, with balanced accuracy from {fmt_num(stats.get('arch_min_ba'))} to {fmt_num(stats.get('arch_max_ba'))}. The maximum permutation p-value across the residualized overall architecture rows was {fmt_p(stats.get('arch_max_perm_p'))}.",
                f"The family gate supports the main mechanistic distinction: state-like operating perturbations and capacity-like representational perturbations leave different behavioural/dynamical fingerprints even when mean performance is similar. However, hybrid scalar recovery is weaker after residualisation. Median raw hybrid state recovery was rho = {fmt_num(stats.get('hybrid_raw_state_median_rho'))}, whereas median residualized hybrid state recovery was rho = {fmt_num(stats.get('hybrid_resid_state_median_rho'))}. This is why the final claim audit grades scalar state recovery as qualified.",
                "For writing, the allowed conclusion is that perturbation families are separable and architecture-robust. The disallowed conclusion is that the analysis has proved a clean architecture-free scalar state axis in humans.",
            ],
        ),
        (
            "Human State Results",
            [
                f"State behaves as a session-task profile rather than a stable trait. The median state ICC is {fmt_num(stats.get('state_median_icc'))}, while capacity ICC is {fmt_num(stats.get('capacity_median_icc'))}. The lower state ICC is expected: it indicates that state has meaningful within-person/session-task variation. It does not mean the state estimate is useless; it means the construct should be interpreted as operating reliability/instability.",
                "State reliability analyses show that trial count matters. Short trial bins are less stable, while larger bins produce better split-half and bootstrap behaviour. This supports the practical recommendation that future validation datasets should include enough trials per session/task before asking state to behave like a robust profile.",
                "The strongest state interpretation is behavioural: state tracks lapse-like error, response variability, drift and reliability. The weakest state interpretation is scalar/latent: the artificial hybrid residualized recovery is not strong enough to justify treating state as a single universal coordinate.",
            ],
        ),
        (
            "Human Capacity Results",
            [
                "Capacity is the strongest construct in the current project. It is participant-level in the current implementation, shows high between-person stability by design and analysis, and remains supported under multiple ablated variants. The capacity variants include full, behaviour-only, model-only, geometry-blind and feature-ablation definitions; all remained at least moderate in the current validation summary.",
                f"The strongest human capacity validation is the TU Berlin load-pressure model. The state-controlled load-by-capacity interaction predicts accuracy with beta = {fmt_num(stats.get('tu_accuracy_load_x_capacity_beta'))}, q = {fmt_p(stats.get('tu_accuracy_load_x_capacity_q'))}, and median reaction time with beta = {fmt_num(stats.get('tu_rt_load_x_capacity_beta'))}, q = {fmt_p(stats.get('tu_rt_load_x_capacity_q'))}. This supports capacity as a resource/pressure profile: load exposes capacity differences rather than merely lowering everyone equally.",
                "Capacity can be written more strongly than state, but still not as a direct neural resource. The safe wording is recurrent-model capacity profile, recurrent-geometry profile, or capacity-pressure profile. The unsafe wording is direct biological capacity measure.",
            ],
        ),
        (
            "Prediction, Incremental Value and Controls",
            [
                f"Prediction analyses show explanatory value but not universal predictive supremacy. In the all-supervised context, additive state-capacity RMSE is {fmt_num(stats.get('rmse_additive_state_capacity'))}, state-capacity interaction RMSE is {fmt_num(stats.get('rmse_state_capacity_interaction'))}, random-axis control RMSE is {fmt_num(stats.get('rmse_random_axis_control'))}, and shuffled-coordinate control RMSE is {fmt_num(stats.get('rmse_shuffled_coordinate_control'))}. This pattern supports real signal in the profiles, because shuffled/random controls do not reproduce the gains.",
                "At the same time, behavioural descriptive baselines remain important. The project should not claim that recurrent profiles always beat simple behavioural summaries. The more defensible claim is that the profiles provide a mechanistically organized decomposition and useful explanatory axes, especially for capacity pressure and state reliability.",
            ],
        ),
        (
            "Physiology Results",
            [
                f"The physiology layer now uses row-level EEG/fNIRS/ECG feature tests and blocked permutation controls rather than simply copying earlier summary outputs. It evaluates {stats.get('physiology_tests', 0)} feature-profile associations; {stats.get('physiology_perm_fdr_rows', 0)} survive permutation FDR. EEG contributes most of the surviving rows, with smaller ECG and fNIRS contributions.",
                "These physiology findings are valuable because they show the profiles are not purely arbitrary behavioural scores. But they are exploratory for three reasons: features are summary physiological markers rather than full encoding models; ds007554 correctness was reconstructed; and the project does not yet include a prospective physiology dataset designed around this exact hypothesis.",
                "The correct manuscript language is therefore bounded physiological alignment. The report, claim audit and manuscript should avoid neural basis, neural coordinate or direct neurophysiological implementation unless future work supplies stronger evidence.",
            ],
        ),
        (
            "Claim Audit and Final Strength Labels",
            [
                f"The final claim audit contains {stats.get('claims_strong', 0)} strong, {stats.get('claims_moderate', 0)} moderate, {stats.get('claims_qualified', 0)} qualified and {stats.get('claims_exploratory', 0)} exploratory claim labels. Strong claims are reserved for architecture-family separation and capacity validation. The state scalar claim is qualified because residualized scalar recovery fails or weakens. Physiology is exploratory.",
                "The audit is not cosmetic. It is a guardrail for writing. Every future manuscript sentence should respect the claim label: strong for family separation and capacity pressure, moderate for state reliability and baseline-constrained explanatory value, qualified for scalar state recovery and exploratory for physiology.",
            ],
        ),
    ]


def build_word_report(stats: dict[str, object], manifest: pd.DataFrame) -> Path:
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    run = title.add_run("Full Results Record: TinyRNN State and Capacity")
    run.font.name = PRESET["font"]
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = rgb(PRESET["ink"])
    add_para(doc, f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} from the current NHB revision outputs.", style="Caption")

    for heading, paragraphs in narrative_sections(stats):
        doc.add_heading(heading, level=1)
        for p in paragraphs:
            add_para(doc, p)

    doc.add_heading("Method and Analysis Step Protocol", level=1)
    add_para(
        doc,
        "The analysis was intentionally divided into step folders and reproducible scripts to avoid overload and to make each claim traceable. The table below records the current step logic used by this record package.",
    )
    add_doc_table(
        doc,
        "Table R0 | Step protocol and purpose.",
        pd.DataFrame(STEP_PROTOCOL, columns=["step", "name", "purpose"]),
        ["step", "name", "purpose"],
        [750, 2100, 6510],
        max_rows=30,
    )

    doc.add_heading("Detailed Results by Analysis Family", level=1)
    for heading, paragraphs in detailed_result_sections(stats):
        doc.add_heading(heading, level=2)
        for p in paragraphs:
            add_para(doc, p)

    doc.add_heading("Dataset Inventory", level=1)
    counts = read_table(TABLES / "event_counts_by_dataset.csv")
    add_doc_table(doc, "Table R1 | Event counts by dataset.", counts, ["dataset", "n_included_events", "n_subjects", "n_sessions", "n_tasks", "n_events_with_correct", "n_events_with_rt"], [1800, 1300, 900, 900, 900, 1400, 1300])

    doc.add_heading("Artificial-Agent and Architecture Results", level=1)
    add_para(doc, "The revised Step 02 is no longer a replay-only analysis. It trains true vanilla RNN, GRU and LSTM artificial agents and evaluates performance-matched state-like versus capacity-like perturbation families.")
    arch = read_table(NHB_TABLES / "architecture_perturbation_gate_results.csv")
    add_doc_table(doc, "Table R2 | Residualized architecture gate, overall rows.", arch[(arch.get("task_family", "") == "overall") & (arch.get("feature_set", "") == "residualized_fingerprint")] if not arch.empty else arch, ["model_family", "feature_set", "balanced_accuracy", "auc", "permutation_p", "claim_strength"], [1300, 1850, 1400, 1000, 1350, 1100])
    hybrid = read_table(NHB_TABLES / "architecture_hybrid_recovery_results.csv")
    add_doc_table(doc, "Table R3 | Hybrid recovery results.", hybrid, ["model_family", "analysis_type", "feature_set", "spearman_rho", "nominal_p_value", "pass_gate", "claim_strength"], [1000, 1600, 1500, 1200, 1300, 900, 1100])

    doc.add_heading("State Results", level=1)
    add_para(doc, "State evidence is retained, but as a bounded operating-profile result. Reliability depends on trial count and dataset/task structure. Variance decomposition supports state as more session/task-variable than capacity.")
    state_var = read_table(NHB_TABLES / "state_capacity_variance_decomposition.csv")
    add_doc_table(doc, "Table R4 | State and capacity variance decomposition.", state_var, ["construct", "component", "n_rows", "n_subjects", "within_person_variance", "between_person_variance", "icc", "claim_strength"], [1000, 1900, 900, 900, 1400, 1400, 900, 1000])
    state_boot = read_table(NHB_TABLES / "state_bootstrap_reliability.csv")
    add_doc_table(doc, "Table R5 | State bootstrap reliability, first rows.", state_boot, ["dataset", "task", "participant_id", "split_type", "n_trials", "trial_count_bin", "cosine_similarity", "bootstrap_median"], [1000, 1000, 1500, 1000, 800, 1000, 1200, 1200])

    doc.add_heading("Capacity Results", level=1)
    add_para(doc, "Capacity is the strongest construct in the current package. It survives architecture checks, remains stable by variance decomposition, and shows load-pressure effects in TU Berlin.")
    cap = read_table(NHB_TABLES / "capacity_pressure_models.csv")
    add_doc_table(doc, "Table R6 | Capacity pressure models.", cap, ["outcome", "predictor", "n_rows", "estimate", "std_error", "p_value", "q_value", "claim_strength"], [1300, 1900, 800, 1000, 1000, 1200, 1200, 1100])
    cap_ab = read_table(NHB_TABLES / "capacity_variant_validation_summary.csv")
    add_doc_table(doc, "Table R7 | Capacity variant validation summary.", cap_ab, ["variant", "n_validations", "median_abs_effect", "claim_strength"], [3200, 1200, 1500, 1200])

    doc.add_heading("Prediction, Generalization and Controls", level=1)
    inc = read_table(NHB_TABLES / "incremental_value_model_comparison.csv")
    add_doc_table(doc, "Table R8 | Incremental value and baseline challenge.", inc, ["dataset", "model_name", "RMSE", "delta_RMSE", "claim_strength"], [2200, 2500, 1000, 1200, 1200])
    shuffle = read_table(NHB_TABLES / "profile_shuffle_controls.csv")
    add_doc_table(doc, "Table R9 | Shuffle/falsification controls.", shuffle, list(shuffle.columns[: min(7, len(shuffle.columns))]), [1300] * min(7, max(1, len(shuffle.columns))))

    doc.add_heading("Physiology Results", level=1)
    add_para(doc, "Physiology is now a real row-level screen rather than a copied summary. The table includes task/session/load residualisation and blocked permutation controls. Surviving rows should still be treated as exploratory external alignment.")
    phys = read_table(NHB_TABLES / "physiology_robustness_models.csv")
    add_doc_table(doc, "Table R10 | Top physiology associations by permutation p value.", phys.sort_values("permutation_p_value") if not phys.empty and "permutation_p_value" in phys else phys, ["dataset", "modality", "feature", "predictor", "n_rows", "estimate", "permutation_p_value", "permutation_q_value", "claim_strength"], [1000, 800, 2100, 2100, 800, 900, 1300, 1300, 1200])

    doc.add_heading("Claim Audit", level=1)
    claims = read_table(NHB_AUDIT / "nhb_final_claim_audit.tsv")
    add_doc_table(doc, "Table R11 | Final claim audit.", claims, ["claim_id", "claim_text", "claim_strength", "controls_failed", "allowed_manuscript_language", "forbidden_language"], [650, 3000, 950, 1600, 2200, 1800])

    doc.add_heading("Result File Index", level=1)
    add_para(doc, "The Excel workbook accompanying this report contains a fuller multi-sheet export. The index below records the result files available at generation time.")
    add_doc_table(doc, "Table R12 | Result file manifest excerpt.", manifest, ["kind", "file", "relative_path", "rows", "columns"], [950, 2300, 4200, 750, 750], max_rows=30)

    doc.add_heading("Limitations to Preserve in Future Writing", level=1)
    for item in [
        "State should not be described as a fully validated scalar coordinate; residualized scalar recovery is qualified.",
        "Physiology should not be described as direct neural-coordinate evidence.",
        "ds007554 correctness was reconstructed from push-button timing, so ds007554 physiology claims require bounded language.",
        "Behavioural descriptive baselines remain strong; coordinate models are explanatory, not universally superior predictors.",
        "Git is not available on the current PowerShell PATH, so repository status could not be recorded from this environment.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)
    return REPORT_DOCX


def sheet_name(base: str, used: set[str]) -> str:
    name = re.sub(r"[\[\]\:\*\?\/\\]", "_", base)[:31] or "sheet"
    candidate = name
    i = 1
    while candidate in used:
        suffix = f"_{i}"
        candidate = f"{name[:31-len(suffix)]}{suffix}"
        i += 1
    used.add(candidate)
    return candidate


def write_excel_sheet(writer: pd.ExcelWriter, name: str, df: pd.DataFrame, used: set[str]) -> str:
    sname = sheet_name(name, used)
    out = df.copy()
    if len(out) > MAX_EXCEL_ROWS:
        note = pd.DataFrame({"warning": [f"Original table had {len(out)} rows; Excel sheet is truncated to {MAX_EXCEL_ROWS} rows."], "source_rows": [len(out)]})
        out = pd.concat([note, out.head(MAX_EXCEL_ROWS - len(note))], ignore_index=True, sort=False)
    out.to_excel(writer, sheet_name=sname, index=False)
    return sname


def build_excel_workbook(stats: dict[str, object], manifest: pd.DataFrame) -> Path:
    RECORD_DIR.mkdir(parents=True, exist_ok=True)
    used: set[str] = set()
    with pd.ExcelWriter(EXCEL_OUT, engine="openpyxl") as writer:
        pd.DataFrame([stats]).to_excel(writer, sheet_name="README_KEY_RESULTS", index=False)
        used.add("README_KEY_RESULTS")
        manifest.to_excel(writer, sheet_name="RESULT_FILE_INDEX", index=False)
        used.add("RESULT_FILE_INDEX")
        pd.DataFrame(STEP_PROTOCOL, columns=["step", "name", "purpose"]).to_excel(writer, sheet_name="STEP_PROTOCOL", index=False)
        used.add("STEP_PROTOCOL")

        for label, root in [("NHB", NHB_TABLES), ("BASE", TABLES), ("AUDIT", NHB_AUDIT), ("FANCY", NHB / "fancy_figures" / "source_data")]:
            if not root.exists():
                continue
            for path in sorted(root.glob("*")):
                if path.is_file() and path.suffix.lower() in {".csv", ".tsv"}:
                    df = read_table(path)
                    if df.empty and path.stat().st_size > 0:
                        continue
                    write_excel_sheet(writer, f"{label}_{path.stem}", df, used)

    wb = load_workbook(EXCEL_OUT)
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = Font(bold=True, color="14213D")
            cell.fill = PatternFill("solid", fgColor="F4F6F9")
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        max_col = min(ws.max_column, 25)
        for col_idx in range(1, max_col + 1):
            letter = get_column_letter(col_idx)
            max_len = 10
            for cell in ws[letter][: min(ws.max_row, 200)]:
                if cell.value is not None:
                    max_len = min(max(max_len, len(str(cell.value)) + 2), 45)
            ws.column_dimensions[letter].width = max_len
        for row in ws.iter_rows(min_row=2, max_row=min(ws.max_row, 1000)):
            for cell in row:
                cell.alignment = Alignment(wrap_text=False, vertical="top")
    wb.save(EXCEL_OUT)
    return EXCEL_OUT


def build_markdown_report(stats: dict[str, object], manifest: pd.DataFrame) -> Path:
    text = [
        "# Full Results Record: TinyRNN State and Capacity",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "## Key Results",
    ]
    for heading, paragraphs in narrative_sections(stats):
        text.append(f"### {heading}")
        text.extend(paragraphs)
    text.append("## Step Protocol")
    for step, name, purpose in STEP_PROTOCOL:
        text.append(f"- Step {step}: **{name}** — {purpose}")
    text.append("## Result File Index")
    for _, row in manifest.head(250).iterrows():
        text.append(
            f"- {row.get('kind', '')}: `{row.get('relative_path', '')}` "
            f"({row.get('rows', '')} rows, {row.get('columns', '')} columns, {row.get('bytes', '')} bytes)"
        )
    if len(manifest) > 250:
        text.append(f"- ... {len(manifest) - 250} additional indexed files are available in the Excel workbook.")
    REPORT_MD.write_text("\n\n".join(text), encoding="utf-8")
    return REPORT_MD


def build_python_reference(stats: dict[str, object], manifest: pd.DataFrame) -> Path:
    result_paths = manifest["relative_path"].dropna().astype(str).tolist()
    safe_stats = {}
    for key, value in stats.items():
        if isinstance(value, pd.DataFrame):
            continue
        if isinstance(value, (np.floating, float)) and not np.isfinite(value):
            safe_stats[key] = None
        elif isinstance(value, (np.integer, int)):
            safe_stats[key] = int(value)
        elif isinstance(value, (np.floating, float)):
            safe_stats[key] = float(value)
        else:
            safe_stats[key] = value
    method_text = f'''
Full method and analysis reference for the TinyRNN state-capacity project.

Generated: {datetime.now().isoformat(timespec='seconds')}

Purpose
-------
This single Python file is a record/reference artifact. It documents the full
method and analysis logic, lists the major result files, and exposes helper
functions for reloading tables from the current workspace.

Current empirical interpretation
--------------------------------
1. Artificial-agent perturbation families are separable across true vanilla RNN,
   GRU and LSTM agents after performance residualisation.
2. Capacity is the strongest current construct: it has architecture robustness,
   participant-level stability, recurrent-geometry convergence and TU Berlin
   load-pressure support.
3. State is meaningful but qualified: it behaves as a session/task reliability
   profile, not a fully validated scalar coordinate.
4. Physiology is real but exploratory: EEG/fNIRS/ECG associations were tested
   with row-level blocked permutation controls, but should not be written as
   direct neural-coordinate evidence.
5. Behavioural baselines remain strong, so recurrent profiles should be framed
   as explanatory/mechanistic profiles rather than universally superior
   predictors.

Selected numeric anchors
------------------------
- Included events: {stats.get('total_events')}
- Architecture residualized balanced accuracy range:
  {fmt_num(stats.get('arch_min_ba'))} to {fmt_num(stats.get('arch_max_ba'))}
- Architecture permutation p: {fmt_p(stats.get('arch_max_perm_p'))}
- Median state ICC: {fmt_num(stats.get('state_median_icc'))}
- Median capacity ICC: {fmt_num(stats.get('capacity_median_icc'))}
- TU accuracy load-by-capacity beta/q:
  {fmt_num(stats.get('tu_accuracy_load_x_capacity_beta'))} /
  {fmt_p(stats.get('tu_accuracy_load_x_capacity_q'))}
- Physiology tests / permutation-FDR rows:
  {stats.get('physiology_tests')} / {stats.get('physiology_perm_fdr_rows')}

Step protocol
-------------
''' + "\n".join([f"Step {s}: {n}. {p}" for s, n, p in STEP_PROTOCOL])

    code = f'''from __future__ import annotations

from pathlib import Path
import pandas as pd

ROOT = Path(__file__).resolve().parents[3]
TABLES = ROOT / "outputs" / "tables"
NHB = ROOT / "outputs" / "nhb_revision"
NHB_TABLES = NHB / "tables"
NHB_AUDIT = NHB / "audit"

METHOD_AND_ANALYSIS_REFERENCE = {method_text!r}

KEY_RESULTS = {json.dumps(safe_stats, indent=2, default=str)}

RESULT_FILES = {json.dumps(result_paths, indent=2)}

STEP_PROTOCOL = {json.dumps([{"step": s, "name": n, "purpose": p} for s, n, p in STEP_PROTOCOL], indent=2)}


def load_table(relative_path: str) -> pd.DataFrame:
    """Load a CSV/TSV result table by relative path from project root."""
    path = ROOT / relative_path
    sep = "\\t" if path.suffix.lower() == ".tsv" else ","
    return pd.read_csv(path, sep=sep)


def load_named_table(name: str, nhb: bool = True) -> pd.DataFrame:
    """Load a table from outputs/nhb_revision/tables or outputs/tables."""
    root = NHB_TABLES if nhb else TABLES
    return load_table(str((root / name).relative_to(ROOT)))


def list_result_files() -> list[str]:
    """Return all result files captured when this reference was generated."""
    return list(RESULT_FILES)


def print_summary() -> None:
    """Print the embedded method/analysis reference narrative."""
    print(METHOD_AND_ANALYSIS_REFERENCE)


if __name__ == "__main__":
    print_summary()
'''
    PY_REF.write_text(code, encoding="utf-8")
    return PY_REF


def build_manifest(outputs: Iterable[Path], stats: dict[str, object], result_manifest: pd.DataFrame) -> Path:
    payload = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "analysis_id": ANALYSIS_ID,
        "script_name": SCRIPT_NAME,
        "outputs": [rel(p) for p in outputs],
        "key_stats": {k: (None if isinstance(v, float) and not np.isfinite(v) else v) for k, v in stats.items() if not isinstance(v, pd.DataFrame)},
        "n_result_files_indexed": int(len(result_manifest)),
    }
    MANIFEST_JSON.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
    return MANIFEST_JSON


def main() -> None:
    ensure_nhb_dirs()
    RECORD_DIR.mkdir(parents=True, exist_ok=True)
    started = datetime.now(timezone.utc).isoformat()
    manifest = collect_result_files()
    stats = key_stats()
    outputs: list[Path] = []
    outputs.append(build_word_report(stats, manifest))
    outputs.append(build_excel_workbook(stats, manifest))
    outputs.append(build_markdown_report(stats, manifest))
    outputs.append(build_python_reference(stats, manifest))
    outputs.append(build_manifest(outputs, stats, manifest))
    append_manifest(ANALYSIS_ID, SCRIPT_NAME, outputs)
    append_registry(ANALYSIS_ID, SCRIPT_NAME, started, outputs, notes="Built record package: detailed Word report, all-results Excel workbook, Markdown report and Python reference.")
    print(f"Wrote record package to {RECORD_DIR}")


if __name__ == "__main__":
    main()
