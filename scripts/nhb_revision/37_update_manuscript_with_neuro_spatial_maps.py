from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
INPUT = Path("C:/Users/Gebruiker/Downloads/State_capacity_MERGED_tracked.docx")
OUT_DIR = ROOT / "outputs" / "nhb_revision" / "neuro_spatial_maps"
FIG = OUT_DIR / "fig_neuro_spatial_eeg_topomaps_nominal_marked.png"
OUTPUT = OUT_DIR / "State_capacity_MERGED_tracked_neuro_spatial_updated.docx"
CHANGE_LOG = OUT_DIR / "manuscript_neuro_spatial_update_log.md"


RESULT_INSERT = (
    "We then extended the physiology screen from global summary features to EEG sensor-space maps. "
    "For ds007554, 515 EEG recordings yielded 65,920 channel-band observations across 32 standard 10-05 sensors. "
    "Channel-wise models tested log band power against state and capacity profiles while adjusting for task and session. "
    "The strongest nominal effects were state-related and concentrated over frontal and temporal sensors (for example, "
    "state-beta at FP1: beta = 0.264, p = 0.0189; state-beta at TP8: beta = 0.233, p = 0.0198), but no sensor survived "
    "FDR correction (minimum q = 0.198). These maps therefore provide a spatial exploratory visualization of the physiology "
    "screen, not corrected evidence for a localized neural substrate of state or capacity (Extended Data Fig. 2)."
)

METHOD_INSERT = (
    "For the EEG spatial visualization, ds007554 EDF recordings were reprocessed at the channel level. "
    "Band power was estimated for delta, theta, alpha and beta bands at each standard 10-05 channel. "
    "Session-level state coordinates and participant-level capacity coordinates were attached using the same repaired "
    "coordinate tables as the physiology screen. For each channel and band, log band power was modelled as a function of "
    "state and capacity profiles with task and session covariates and HC3 robust standard errors. Channel-wise p values "
    "were FDR corrected across the sensor-space family. Topographic maps show task/session-adjusted coefficients; open "
    "circles indicate nominal p < 0.05 only, because no sensor survived FDR q < 0.05."
)

EXT_FIG_CAPTION = (
    "Extended Data Fig. 2 | EEG sensor-space state-capacity maps. Topographic maps show task/session-adjusted coefficients "
    "from channel-wise ds007554 EEG models for theta, alpha and beta band power. Open circles mark nominal p < 0.05 sensors; "
    "no sensor survived FDR q < 0.05, so the maps are exploratory visualization rather than corrected source-localization "
    "evidence. Source data are provided in the Step 35 neuro-spatial workbook."
)


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    p.add_run(text)
    return p


def insert_picture_after(paragraph, image_path: Path, width_inches: float = 6.4):
    p = insert_paragraph_after(paragraph, "")
    p.alignment = 1
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def replace_exact(doc: Document, old: str, new: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text == old:
            paragraph.text = new
            return True
    return False


def find_exact(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    return None


def integrity(path: Path) -> dict[str, int | str | None]:
    result: dict[str, int | str | None] = {}
    with zipfile.ZipFile(path) as zf:
        result["bad_member"] = zf.testzip()
        names = zf.namelist()
        result["duplicate_members"] = len(names) - len(set(names))
        for xml_name in ["[Content_Types].xml", "word/document.xml", "word/settings.xml", "word/_rels/document.xml.rels"]:
            etree.fromstring(zf.read(xml_name))
        result["media_count"] = len([n for n in names if n.startswith("word/media/")])
    d = Document(str(path))
    result["paragraph_count"] = len(d.paragraphs)
    result["inline_shapes"] = len(d.inline_shapes)
    return result


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    if not FIG.exists():
        raise FileNotFoundError(FIG)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(str(INPUT))

    # Results: insert after the existing bounded physiology paragraph.
    physiology_boundary = find_exact(
        doc,
        "These physiological associations are useful as external convergence, but they are not neural coordinates. The features are summary markers rather than encoding-model estimates; ds007554 behavioural correctness was partly reconstructed; and no dataset was prospectively designed for this hypothesis. The physiology claim is therefore labelled exploratory in the claim audit.",
    )
    if physiology_boundary is None:
        raise RuntimeError("Could not locate physiology boundary paragraph")
    insert_paragraph_after(physiology_boundary, RESULT_INSERT)

    # Methods: extend physiology screen methods with the channel-wise map analysis.
    methods_anchor = find_exact(
        doc,
        "Physiology analyses used summary features, not encoding models. EEG features included frontal theta power, relative alpha/theta measures, theta-to-alpha ratio, aperiodic slope and channel variability53-57. fNIRS features included variability of oxygenated and deoxygenated haemoglobin channels. ECG features included RMSSD and SDNN58-60. Associations with state and capacity profiles were tested after task, session and load residualisation. Null distributions were generated by subject-blocked permutation, and multiplicity was controlled by permutation-based FDR. Surviving physiology results are reported as exploratory external alignment only.",
    )
    if methods_anchor is None:
        raise RuntimeError("Could not locate physiology methods paragraph")
    insert_paragraph_after(methods_anchor, METHOD_INSERT)

    # Add Extended Data figure after the main physiology figure caption, before Discussion.
    fig6_caption = None
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Fig. 6 | Exploratory physiology"):
            fig6_caption = paragraph
            break
    if fig6_caption is None:
        raise RuntimeError("Could not locate Fig. 6 caption")
    p = insert_picture_after(fig6_caption, FIG, width_inches=6.4)
    insert_paragraph_after(p, EXT_FIG_CAPTION)

    # Update the Extended Data note to include the new sensor-space map and its source data.
    old_note = (
        "Extended Data Fig. 1 and Extended Data Tables 1-5 provide the model-comparison baseline, artificial architecture and hybrid-recovery tables, load-pressure coefficients, state variance components, capacity ablation/variant checks, physiology associations and the full claim-audit record. These tables are essential for reviewers because the main text deliberately distinguishes profile-level support from failed scalar-coordinate gates."
    )
    new_note = (
        "Extended Data Figs. 1-2 and Extended Data Tables 1-5 provide the model-comparison baseline, artificial architecture and hybrid-recovery tables, load-pressure coefficients, state variance components, capacity ablation/variant checks, physiology associations, EEG sensor-space maps and the full claim-audit record. These materials are essential for reviewers because the main text deliberately distinguishes profile-level support from failed scalar-coordinate gates and exploratory physiology visualizations."
    )
    if not replace_exact(doc, old_note, new_note):
        raise RuntimeError("Could not update Extended Data note")

    doc.save(OUTPUT)
    check = integrity(OUTPUT)
    CHANGE_LOG.write_text(
        "\n".join(
            [
                "# Manuscript neuro-spatial update log",
                "",
                f"Input: `{INPUT}`",
                f"Output: `{OUTPUT}`",
                "",
                "Updates made:",
                "- Added a Results paragraph describing the Step 35 EEG sensor-space analysis.",
                "- Added a Methods paragraph documenting channel-wise EEG band-power models.",
                "- Inserted the nominal-marker EEG topomap as Extended Data Fig. 2 after Fig. 6.",
                "- Updated the Extended Data note to include EEG sensor-space maps.",
                "",
                "Claim boundary preserved:",
                "- Open-circle markers are nominal p < 0.05 only.",
                "- No sensor survived FDR q < 0.05.",
                "- The manuscript text describes the analysis as exploratory visualization, not localized neural evidence.",
                "",
                f"Integrity check: {check}",
            ]
        ),
        encoding="utf-8",
    )
    print(OUTPUT)
    print(CHANGE_LOG)
    print(check)


if __name__ == "__main__":
    main()
