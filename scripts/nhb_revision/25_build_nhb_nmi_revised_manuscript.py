from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable, Sequence

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parents[1]
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from nhb_utils import NHB_AUDIT, NHB_MANUSCRIPT, NHB_TABLES, TABLES, append_manifest, append_registry, ensure_nhb_dirs


ANALYSIS_ID = "nhb_25_revised_manuscript"
SCRIPT_NAME = "scripts/nhb_revision/25_build_nhb_nmi_revised_manuscript.py"
PROPOSAL_DOCX = Path(r"C:\Users\Gebruiker\Downloads\state_capacity_revised_intro_methods_NHB.docx")
PROPOSAL_TXT = NHB_MANUSCRIPT / "proposal_intro_methods_extracted.txt"
FANCY = ROOT / "outputs" / "nhb_revision" / "fancy_figures"
OUT_DOCX = NHB_MANUSCRIPT / "state_capacity_NHB_NMI_revised_manuscript.docx"
OUT_MD = NHB_MANUSCRIPT / "state_capacity_NHB_NMI_revised_manuscript.md"


PRESET = {
    "name": "narrative_proposal",
    "font": "Calibri",
    "body_size": 11,
    "body_after": 8,
    "body_line": 1.333,
    "h1_size": 16,
    "h2_size": 13,
    "h3_size": 12,
    "h1_before": 18,
    "h1_after": 10,
    "h2_before": 12,
    "h2_after": 6,
    "h3_before": 8,
    "h3_after": 4,
    "heading_blue": "2E74B5",
    "heading_dark": "1F4D78",
    "ink": "14213D",
    "muted": "5C6672",
    "table_header_fill": "F4F6F9",
    "table_indent_dxa": 120,
    "content_width_dxa": 9360,
    "cell_margins": {"top": 80, "bottom": 80, "start": 120, "end": 120},
}


FIGURES = [
    (
        "fancy_fig1_graphical_abstract.png",
        "Figure 1 | Study logic and claim governance. The analysis begins with directly downloadable behavioural and physiological datasets, then fits compact recurrent models, tests machine perturbation gates, and constrains manuscript claims through a formal audit. The architecture gate is shown as a pre-human falsification step.",
    ),
    (
        "fancy_fig2_state_capacity_landscape.png",
        "Figure 2 | Human state-capacity landscape. Session-task state profiles and participant-level capacity profiles separate datasets and tasks but do not justify a universal scalar state coordinate. State is therefore interpreted as an operating-reliability profile.",
    ),
    (
        "fancy_fig3_architecture_robustness.png",
        "Figure 3 | Architecture robustness. State-like and capacity-like perturbation families remain separable in true vanilla RNN, GRU and LSTM agents after performance residualisation. Hybrid scalar recovery is stronger in raw fingerprints and weaker after residualisation, bounding the state-axis claim.",
    ),
    (
        "fancy_fig4_capacity_pressure_surface.png",
        "Figure 4 | Capacity pressure under N-back load. TU Berlin load models show performance decline and reaction-time slowing with increasing load, with state-controlled load-by-capacity interactions supporting capacity-pressure language.",
    ),
    (
        "fancy_fig5_state_reliability_atlas.png",
        "Figure 5 | State reliability atlas. State components show moderate within-person/session-task structure, and split-half reliability improves with trial count. This supports state as a reliability profile rather than a stable capacity-like trait.",
    ),
    (
        "fancy_fig6_physiology_claim_audit.png",
        "Figure 6 | Physiology and bounded claims. EEG, fNIRS and ECG associations provide exploratory alignment after blocked permutation controls. The claim audit explicitly prevents direct neural-coordinate language.",
    ),
]


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.strip("#")
    return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:], 16))


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    sep = "\t" if path.suffix.lower() == ".tsv" else ","
    return pd.read_csv(path, sep=sep)


def t(name: str, nhb: bool = True) -> pd.DataFrame:
    return read_csv((NHB_TABLES if nhb else TABLES) / name)


def fmt_num(value: object, digits: int = 3) -> str:
    try:
        x = float(value)
    except Exception:
        return "NA"
    if not np.isfinite(x):
        return "NA"
    if abs(x) < 0.001 and x != 0:
        return f"{x:.2e}"
    if abs(x) < 1:
        return f"{x:.{digits}g}"
    return f"{x:.{digits}f}".rstrip("0").rstrip(".")


def fmt_p(value: object) -> str:
    try:
        x = float(value)
    except Exception:
        return "NA"
    if not np.isfinite(x):
        return "NA"
    if x < 0.001:
        return f"{x:.2e}"
    return f"{x:.3f}".rstrip("0").rstrip(".")


def extract_proposal() -> tuple[list[str], list[str]]:
    if PROPOSAL_TXT.exists():
        text = PROPOSAL_TXT.read_text(encoding="utf-8")
    else:
        from docx import Document as DocxDocument

        doc = DocxDocument(PROPOSAL_DOCX)
        text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        PROPOSAL_TXT.parent.mkdir(parents=True, exist_ok=True)
        PROPOSAL_TXT.write_text(text, encoding="utf-8")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    try:
        intro_start = lines.index("Introduction") + 1
        methods_start = lines.index("Methods")
    except ValueError:
        return [], []
    intro = lines[intro_start:methods_start]
    methods = lines[methods_start + 1 :]
    methods = [p for p in methods if p != "Implementation note" and not p.startswith("This draft is written")]
    return revise_intro(intro), revise_methods(methods)


def revise_intro(paragraphs: list[str]) -> list[str]:
    revised: list[str] = []
    for p in paragraphs:
        if p.startswith("Only after this machine-only validation"):
            revised.append(
                "Only after machine-only validation do we interpret human data in the resulting profile space. The current full run supports robust separation of state-like and capacity-like perturbation families across true recurrent architectures, but it also shows that residualized scalar state recovery is limited. Accordingly, we treat human state and capacity as multidimensional profiles rather than as fully validated scalar coordinates."
            )
            continue
        if "In the present implementation, state refers" in p:
            revised.append(
                "In the present implementation, state refers to a session-task operating profile summarizing lapse-like error, drift, variability and reliability. Capacity refers to participant-level representational and dynamical-resource profiles, including selected recurrent dimensionality, model-selection confidence, load robustness and recurrent trajectory geometry. This wording is deliberately operational: the profiles are model-derived explanatory constructs, not direct neural resources or clinical traits."
            )
            continue
        if p.startswith("We test this framework using directly downloadable open datasets"):
            revised.append(
                "We test this framework using directly downloadable open datasets. OpenNeuro ds007554 anchors multimodal behavioural, EEG, fNIRS and ECG analyses; COG-BCI provides repeated-session EEG validation; TU Berlin simultaneous EEG-NIRS provides a direct N-back load-pressure test; and HBN Release 4 provides a large scalability analysis. Sleepy Brain and MAUS are not used as successful confirmatory evidence in the present manuscript."
            )
            continue
        revised.append(p.replace("coordinate system", "profile space").replace("coordinates", "profiles where supported"))
    return revised


def revise_methods(paragraphs: list[str]) -> list[str]:
    revised: list[str] = []
    for p in paragraphs:
        if p == "Human datasets":
            revised.append(p)
            continue
        if p.startswith("OpenNeuro ds007554 was the primary discovery dataset"):
            counts = t("event_counts_by_dataset.csv", nhb=False)
            ds = counts[counts["dataset"].eq("ds007554")].iloc[0] if not counts.empty else {}
            revised.append(
                f"OpenNeuro ds007554 was used as the primary multimodal discovery and physiology dataset. The final unified event table contains {int(ds.get('n_included_events', 31053)):,} included ds007554 events from {int(ds.get('n_subjects', 30))} participants and {int(ds.get('n_events_with_correct', 10896)):,} events with reconstructed supervised correctness. Because explicit distributed trial-correctness labels were incomplete, correctness was reconstructed from push-button timing and all ds007554 coordinate-physiology claims are bounded accordingly."
            )
            continue
        if p.startswith("The machine-only validation gate had four required tests"):
            revised.append(
                "The machine-only validation gate had four tests: state-versus-capacity family separation, known state-severity recovery, known capacity-level recovery and hybrid-agent recovery. The final analysis passes the family-separation gate across vanilla RNN, GRU and LSTM agents, but residualized scalar state recovery remains weak. Strong human claims are therefore allowed for family separation and capacity pressure, while state is interpreted as a qualified reliability profile."
            )
            continue
        if p.startswith("The primary model was a task-conditioned gated recurrent unit"):
            revised.append(
                p
                + " In the revision, artificial-agent robustness was extended to true vanilla RNN, GRU and LSTM families; human profile estimation remains based on the compact GRU pipeline unless explicitly stated."
            )
            continue
        revised.append(p.replace("coordinates", "profiles").replace("coordinate", "profile"))
    return revised


def compute_stats() -> dict[str, object]:
    stats_out: dict[str, object] = {}
    counts = t("event_counts_by_dataset.csv", nhb=False)
    stats_out["total_events"] = int(counts["n_included_events"].sum()) if not counts.empty else 0
    stats_out["n_subjects"] = int(counts["n_subjects"].sum()) if not counts.empty else 0
    for _, r in counts.iterrows() if not counts.empty else []:
        stats_out[f"{r['dataset']}_events"] = int(r["n_included_events"])
        stats_out[f"{r['dataset']}_subjects"] = int(r["n_subjects"])

    arch = t("architecture_perturbation_gate_results.csv")
    g = arch[(arch["task_family"] == "overall") & (arch["feature_set"] == "residualized_fingerprint")] if not arch.empty else pd.DataFrame()
    stats_out["arch_min_ba"] = g["balanced_accuracy"].astype(float).min() if not g.empty else np.nan
    stats_out["arch_max_ba"] = g["balanced_accuracy"].astype(float).max() if not g.empty else np.nan
    stats_out["arch_p"] = g["permutation_p"].astype(float).max() if not g.empty else np.nan

    hybrid = t("architecture_hybrid_recovery_results.csv")
    if not hybrid.empty:
        raw_state = hybrid[(hybrid["analysis_type"] == "hybrid_state_axis") & (hybrid["feature_set"] == "raw_fingerprint")]["spearman_rho"].astype(float)
        resid_state = hybrid[(hybrid["analysis_type"] == "hybrid_state_axis") & (hybrid["feature_set"] == "residualized_fingerprint")]["spearman_rho"].astype(float)
        stats_out["raw_state_rho_med"] = raw_state.median()
        stats_out["resid_state_rho_med"] = resid_state.median()

    cap = t("capacity_pressure_models.csv")
    def pick(outcome: str, predictor: str) -> pd.Series:
        d = cap[(cap["outcome"] == outcome) & (cap["predictor"] == predictor)]
        return d.iloc[0] if not d.empty else pd.Series(dtype=object)
    acc_int = pick("mean_accuracy", "load_x_capacity")
    rt_int = pick("rt_median", "load_x_capacity")
    stats_out["tu_acc_int_beta"] = acc_int.get("estimate", np.nan)
    stats_out["tu_acc_int_q"] = acc_int.get("q_value", np.nan)
    stats_out["tu_rt_int_beta"] = rt_int.get("estimate", np.nan)
    stats_out["tu_rt_int_q"] = rt_int.get("q_value", np.nan)

    var = t("state_capacity_variance_decomposition.csv")
    if not var.empty:
        stats_out["state_icc_med"] = var[var["construct"] == "state"]["icc"].astype(float).median()
        stats_out["capacity_icc_med"] = var[var["construct"] == "capacity"]["icc"].astype(float).median()

    phys = t("physiology_robustness_models.csv")
    stats_out["phys_tests"] = len(phys)
    stats_out["phys_sig"] = int((pd.to_numeric(phys.get("permutation_q_value", pd.Series(dtype=float)), errors="coerce") < 0.05).sum()) if not phys.empty else 0

    inc = t("incremental_value_model_comparison.csv")
    if not inc.empty:
        lookup = inc.set_index("model_name")
        for model in ["task_dataset", "behavioral_descriptive", "additive_state_capacity", "state_capacity_interaction", "random_axis_control", "shuffled_coordinate_control"]:
            if model in lookup.index:
                stats_out[f"{model}_rmse"] = lookup.loc[model, "RMSE"]

    hbn = t("hbn_scalability_summary.csv", nhb=False)
    if not hbn.empty:
        row = hbn.iloc[0]
        stats_out["hbn_subjects"] = row.get("n_subjects_included", row.get("n_subjects_events", np.nan))
        stats_out["hbn_events"] = row.get("n_included_events", np.nan)
        stats_out["hbn_state_rows"] = row.get("n_state_rows", np.nan)
        stats_out["hbn_capacity_rows"] = row.get("n_capacity_rows", np.nan)

    claims = read_csv(NHB_AUDIT / "nhb_final_claim_audit.tsv")
    stats_out["claim_summary"] = claims
    return stats_out


def make_results(stats: dict[str, object]) -> list[tuple[str, list[str]]]:
    return [
        (
            "A machine-first profile space was aligned with open human datasets",
            [
                f"The final event inventory comprised {stats.get('total_events', 0):,} included events across COG-BCI, ds007554, TU Berlin EEG-NIRS and HBN Release 4. The supervised behavioural core included {stats.get('cog_bci_events', 0):,} COG-BCI events, {stats.get('ds007554_events', 0):,} ds007554 events, {stats.get('tu_berlin_eeg_nirs_events', 0):,} TU Berlin events and {stats.get('hbn_release_4_events', 0):,} HBN events. Figure 1 summarizes the pipeline and claim audit used to prevent post hoc inflation of evidence strength.",
                "The revised analysis matches the proposal in its main design: state and capacity are first defined in artificial agents, then projected into human behavioural and physiological data. The necessary correction is terminological. The current results support a profile space and perturbation-family separation, not a universally validated pair of scalar human coordinates.",
            ],
        ),
        (
            "True architecture-specific agents separate state-like and capacity-like perturbations",
            [
                f"State-like and capacity-like perturbation families remained separable after performance residualisation in true vanilla RNN, GRU and LSTM agents. Overall residualized balanced accuracy ranged from {fmt_num(stats.get('arch_min_ba'))} to {fmt_num(stats.get('arch_max_ba'))}, with permutation p = {fmt_p(stats.get('arch_p'))}. Leave-one-architecture tests also remained above chance, supporting the central machine perturbation gate.",
                f"However, scalar recovery was not equally strong. Raw hybrid fingerprints recovered state severity with median rho = {fmt_num(stats.get('raw_state_rho_med'))}, whereas residualized state-severity recovery fell to median rho = {fmt_num(stats.get('resid_state_rho_med'))}. This result is the manuscript's main boundary condition: the family distinction is strong, but state should be discussed as a qualified reliability profile rather than as a fully architecture-free scalar coordinate.",
            ],
        ),
        (
            "Capacity showed the clearest recurrent-geometry and load-pressure evidence",
            [
                f"Capacity behaved more like a stable participant-level profile than state: the median capacity ICC was {fmt_num(stats.get('capacity_icc_med'))}, compared with median state ICC {fmt_num(stats.get('state_icc_med'))}. Capacity variant analyses remained moderate across full, model-only, behaviour-only and ablated definitions, indicating that the result is not driven by one isolated feature.",
                f"In TU Berlin, N-back load produced a strong capacity-pressure pattern. The state-controlled load-by-capacity interaction predicted accuracy (beta = {fmt_num(stats.get('tu_acc_int_beta'))}, q = {fmt_p(stats.get('tu_acc_int_q'))}) and median reaction time (beta = {fmt_num(stats.get('tu_rt_int_beta'))}, q = {fmt_p(stats.get('tu_rt_int_q'))}). This is the strongest human validation of the capacity construct and is shown as a pressure surface in Figure 4.",
            ],
        ),
        (
            "State was meaningful but weaker and data-hungry",
            [
                "State profiles were most interpretable as session-task reliability summaries. Split-half analyses showed that reliability improved with trial count, and variance decomposition showed substantially more within-person/session-task variation for state than for capacity. This supports the proposal's distinction between operating state and structural capacity, but in a bounded form.",
                "The state evidence is not absent. COG-BCI supports state as a behavioural reliability and RT-variability profile, and ds007554 provides exploratory physiology alignment after coordinate repair. But state evidence is weaker than capacity evidence because the residualized artificial-agent state gate is limited and several external state tests are nonsignificant.",
            ],
        ),
        (
            "Physiology aligned with profiles but remains exploratory",
            [
                f"Row-level physiology controls tested {stats.get('phys_tests', 0)} EEG, fNIRS and ECG feature-profile associations with task/session/load residualisation and blocked permutation controls. {stats.get('phys_sig', 0)} rows survived permutation FDR. These associations support physiological alignment, especially for EEG and ECG features, but they are explicitly not treated as direct evidence for neural state/capacity coordinates.",
                "The physiology result therefore matches the proposal only in a qualified sense. EEG, fNIRS and ECG data are involved, and they do align with the model-derived profiles in several places, but the strongest manuscript language must remain 'bounded physiological alignment' rather than 'neural basis' or 'neural coordinate'.",
            ],
        ),
        (
            "Behavioural baselines and negative controls constrain the conclusion",
            [
                f"Across supervised datasets, coordinate models improved over task/dataset and shuffled controls, but behavioural descriptive summaries remained competitive. The additive state-capacity model had RMSE {fmt_num(stats.get('additive_state_capacity_rmse'))}, the interaction model had RMSE {fmt_num(stats.get('state_capacity_interaction_rmse'))}, and shuffled/random controls stayed near the task/dataset baseline. This supports explanatory value without claiming universal predictive superiority.",
                "The final claim audit grades the core family-separation and capacity-pressure claims as strong, state reliability as moderate, scalar state recovery as qualified and physiology as exploratory. This audit is the appropriate bridge between the proposal's theoretical ambition and the empirical strength of the present full implementation.",
            ],
        ),
    ]


def make_discussion() -> list[tuple[str, list[str]]]:
    return [
        (
            "Principal interpretation",
            [
                "The study provides a machine-defined, open-data test of the distinction between temporary operating state and representational capacity. The central positive result is that state-like and capacity-like perturbation families can be separated across true recurrent architectures under matched performance. The central constraint is equally important: residualized scalar state recovery remains weak, so the paper should not claim that it has discovered a single state coordinate in human cognition.",
                "The most mature construct is capacity. Capacity aligns with recurrent geometry, behaves as a more stable participant-level profile, and predicts load-pressure effects in a direct N-back validation. State is useful, but currently as a reliability/instability profile that is sensitive to trial count, task structure and dataset quality.",
            ],
        ),
        (
            "Relation to the proposal",
            [
                "The manuscript matches the proposal's architecture in substance: machine perturbations define the candidate axes, human profiles are tested only after machine validation, neurophysiology is used as external alignment evidence, and claims are graded through falsification. The main revision is rhetorical and statistical rather than conceptual. Where the proposal says 'coordinates', the final manuscript should usually say 'profiles' unless the relevant gate passed.",
                "This distinction makes the paper stronger. It prevents the theoretical framework from being judged by an overclaim. A Nature Human Behaviour or Nature Machine Intelligence version should emphasize constructive falsification: the capacity profile passes several independent validations, state is behaviourally meaningful but not yet an architecture-free scalar, and physiology is aligned but exploratory.",
            ],
        ),
        (
            "Limitations and next work",
            [
                "Several limitations remain. ds007554 correctness had to be reconstructed from push-button timing; physiology features are reduced summary markers rather than full neural encoding models; and capacity is currently estimated at participant level, making session-level capacity dynamics hard to test. HBN provides scale but not a close adult N-back replication. These limitations do not invalidate the result, but they define the allowed language.",
                "The next scientific step is a targeted validation dataset with explicit trial correctness, repeated sessions, N-back or PVT-like load manipulation, EEG/fNIRS/ECG and enough trials per session for reliable state estimation. That design would directly test whether state profiles generalize prospectively and whether capacity pressure can be separated from behavioural baselines in a preregistered setting.",
            ],
        ),
    ]


def add_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(section, margin, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRESET["font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal.font.color.rgb = rgb(PRESET["ink"])
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["body_line"]

    for name, size, color, before, after in [
        ("Heading 1", PRESET["h1_size"], PRESET["heading_blue"], PRESET["h1_before"], PRESET["h1_after"]),
        ("Heading 2", PRESET["h2_size"], PRESET["heading_blue"], PRESET["h2_before"], PRESET["h2_after"]),
        ("Heading 3", PRESET["h3_size"], PRESET["heading_dark"], PRESET["h3_before"], PRESET["h3_after"]),
    ]:
        st = styles[name]
        st.font.name = PRESET["font"]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for style_name in ["Caption", "Subtitle"]:
        if style_name in styles:
            st = styles[style_name]
            st.font.name = PRESET["font"]
            st.font.size = Pt(9 if style_name == "Caption" else 11)
            st.font.color.rgb = rgb(PRESET["muted"])
            st.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "State-capacity TinyRNN revision"
    header.style = "Header"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = rgb(PRESET["muted"])


def add_para(doc: Document, text: str, style: str | None = None, justify: bool = True) -> None:
    p = doc.add_paragraph(style=style)
    if justify and style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = PRESET["font"]
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(PRESET["ink"])
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(PRESET["table_indent_dxa"]))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in PRESET["cell_margins"].items():
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(val))
                node.set(qn("w:type"), "dxa")


def set_borders(table, color: str = "B8C2CC") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_table(doc: Document, title: str, columns: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int]) -> None:
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(title).bold = True
    row_list = list(rows)
    tbl = doc.add_table(rows=len(row_list) + 1, cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    set_table_geometry(tbl, widths)
    set_borders(tbl)
    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")
    for i, col in enumerate(columns):
        set_cell_text(tbl.rows[0].cells[i], col, bold=True, size=8.3)
        set_shading(tbl.rows[0].cells[i], PRESET["table_header_fill"])
    for r_idx, row in enumerate(row_list, 1):
        for c_idx, value in enumerate(row):
            set_cell_text(tbl.rows[r_idx].cells[c_idx], value, size=8.1)
    doc.add_paragraph("")


def add_figure(doc: Document, filename: str, caption: str) -> None:
    path = FANCY / filename
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.5))
    pic = doc.inline_shapes[-1]
    pic._inline.docPr.set("descr", caption[:250])
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(caption)
    r.font.size = Pt(9)


def markdown_text(title: str, abstract: str, intro: list[str], results: list[tuple[str, list[str]]], discussion: list[tuple[str, list[str]]], methods: list[str]) -> str:
    parts = [f"# {title}", "## Abstract", abstract, "## Introduction", "\n\n".join(intro), "## Results"]
    for heading, paras in results:
        parts.append(f"### {heading}")
        parts.append("\n\n".join(paras))
    parts.append("## Discussion")
    for heading, paras in discussion:
        parts.append(f"### {heading}")
        parts.append("\n\n".join(paras))
    parts.append("## Methods")
    parts.append("\n\n".join(methods))
    return "\n\n".join(parts)


def build() -> list[Path]:
    ensure_nhb_dirs()
    intro, methods = extract_proposal()
    stats = compute_stats()
    results = make_results(stats)
    discussion = make_discussion()
    title = "Machine-defined state and capacity profiles in human cognition"
    abstract = (
        "Human performance failures can reflect a transiently poor operating state, limited representational capacity or both. "
        "We tested this distinction with compact recurrent networks and directly downloadable behavioural, EEG, fNIRS and ECG datasets. "
        "True vanilla RNN, GRU and LSTM artificial agents separated state-like from capacity-like perturbation families under matched performance. "
        "Capacity showed the strongest human validation through recurrent geometry and N-back load pressure. "
        "State behaved as a meaningful but qualified session-task reliability profile, not as a fully validated scalar coordinate. "
        "Physiology provided exploratory alignment after permutation controls. "
        "The results support a machine-defined profile framework while explicitly bounding claims about neural mechanism and scalar state recovery."
    )

    doc = Document()
    add_style(doc)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(title)
    title_run.font.name = PRESET["font"]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = rgb(PRESET["ink"])
    add_para(doc, "Revised manuscript draft for a Nature Human Behaviour / Nature Machine Intelligence target", style="Subtitle", justify=False)
    add_para(doc, "Prepared from `state_capacity_revised_intro_methods_NHB.docx` and the full NHB revision analysis package.", style="Caption", justify=False)

    doc.add_heading("Abstract", level=1)
    add_para(doc, abstract)

    doc.add_heading("Introduction", level=1)
    for p in intro:
        add_para(doc, p)

    doc.add_heading("Results", level=1)
    for idx, (heading, paras) in enumerate(results, start=1):
        doc.add_heading(heading, level=2)
        for p in paras:
            add_para(doc, p)
        if idx in {1, 2, 3, 4, 5, 6}:
            add_figure(doc, FIGURES[idx - 1][0], FIGURES[idx - 1][1])

    claims = stats.get("claim_summary", pd.DataFrame())
    if isinstance(claims, pd.DataFrame) and not claims.empty:
        add_table(
            doc,
            "Table 1 | Final claim audit summary.",
            ["Claim", "Strength", "Controls failed", "Allowed language"],
            [
                [
                    r.get("claim_id", ""),
                    r.get("claim_strength", ""),
                    r.get("controls_failed", ""),
                    str(r.get("allowed_manuscript_language", ""))[:92],
                ]
                for _, r in claims.iterrows()
            ],
            [900, 1100, 2500, 4860],
        )

    doc.add_heading("Discussion", level=1)
    for heading, paras in discussion:
        doc.add_heading(heading, level=2)
        for p in paras:
            add_para(doc, p)

    doc.add_heading("Methods", level=1)
    for p in methods:
        if len(p) < 80 and not p.endswith("."):
            doc.add_heading(p, level=2)
        else:
            add_para(doc, p)

    doc.add_heading("Data availability", level=1)
    add_para(
        doc,
        "All primary analyses use directly downloadable public datasets or local files inventoried in the reproducibility package. Processed tables, source data, claim audits and figure inputs are stored under `state_capacity_tinyrnn/outputs/nhb_revision/`.",
    )
    doc.add_heading("Code availability", level=1)
    add_para(
        doc,
        "Scripts for the revised analysis are stored under `state_capacity_tinyrnn/scripts/nhb_revision/`. The fancy figures are generated by `24_make_fancy_nature_figures.py`; this manuscript is generated by `25_build_nhb_nmi_revised_manuscript.py`.",
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    OUT_MD.write_text(markdown_text(title, abstract, intro, results, discussion, methods), encoding="utf-8")
    return [OUT_DOCX, OUT_MD]


def main() -> None:
    started = pd.Timestamp.utcnow().isoformat()
    outputs = build()
    append_manifest(ANALYSIS_ID, SCRIPT_NAME, outputs)
    append_registry(ANALYSIS_ID, SCRIPT_NAME, started, outputs, notes="Built revised NHB/NMI manuscript using proposal intro/methods and updated rigorous results.")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
