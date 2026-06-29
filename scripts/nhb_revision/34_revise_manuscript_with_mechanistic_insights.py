from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
INPUT = Path("C:/Users/Gebruiker/Downloads/State_capacity_Draft1_PW.docx")
OUT_DIR = ROOT / "outputs" / "nhb_revision" / "mechanistic_insights"
FIG = OUT_DIR / "fig_mechanistic_insights.png"
OUTPUT = OUT_DIR / "State_capacity_Draft1_PW_mechanistic_revision.docx"
CHANGE_LOG = OUT_DIR / "manuscript_mechanistic_revision_log.md"


NEW_RESULTS_HEADING = "Mechanistic pathways link perturbation family, recurrent dynamics and behaviour"
NEW_RESULTS_PARAS = [
    (
        "The family-separation result shows that graph-preserving and graph-altering perturbations are distinguishable, "
        "but it does not by itself identify the mechanism carrying that boundary. We therefore added three mechanistic "
        "analyses. First, in artificial agents, where perturbation family is known by construction, we tested whether "
        "recurrent-dynamical features mediated the link between imposed family and behaviour. Five of 30 family-to-behaviour "
        "pathways passed the bootstrap/FDR mechanistic screen. The strongest indirect paths ran through recurrent rank, "
        "trajectory radius and hidden-step dynamics, linking the imposed perturbation family to negative log likelihood, "
        "response entropy and probability volatility."
    ),
    (
        "Second, we asked which fingerprint components carried the state-capacity family boundary. A grouped cross-validated "
        "classifier using behavioural and dynamical fingerprint features achieved balanced accuracy of 0.926 and AUC of 0.960. "
        "Permutation attribution concentrated on recurrent-rank dynamics, with a mean importance of 0.333, while the summed "
        "importance of dynamical features exceeded that of behavioural features (0.393 versus 0.076). Thus the main machine "
        "boundary is not simply a behavioural-performance boundary; it is carried primarily by the organisation of recurrent "
        "state space."
    ),
    (
        "Third, perturbation-response contrasts supplied a counterfactual check. Graph-altering capacity scale increased recurrent "
        "rank by a median of 2.5 units and increased trajectory participation ratio by a median of 0.42, whereas graph-preserving "
        "state interventions primarily changed operating-regime variables such as response entropy, hidden variability and probability "
        "volatility. In the human data, observational pathway tests linked capacity and state profiles to accuracy and lapses through "
        "fitted recurrent capacity geometry and trajectory participation ratio. These human paths are not interpreted causally, but they "
        "show that the same fitted dynamical quantities that carry the artificial perturbation boundary also account for profile-behaviour "
        "covariance in people."
    ),
]
NEW_FIG_CAPTION = (
    "Fig. 4 | Mechanistic pathway analyses. a, Artificial-agent mediation screen testing whether imposed perturbation family affects "
    "behaviour through recurrent-dynamical mediators. Asterisks mark bootstrap/FDR-supported indirect paths. b, Supported human "
    "observational pathways. C, capacity profile; S, state-instability profile; G, fitted recurrent capacity geometry; PR, trajectory "
    "participation ratio; Acc, mean accuracy. c, Permutation attribution for the state-capacity family boundary, showing that recurrent "
    "rank dynamics dominate cross-validated classification. d, Perturbation-response contrasts: graph-altering capacity scale shifts "
    "rank and trajectory geometry, whereas graph-preserving state interventions primarily alter operating-regime variables. Human pathways "
    "are observational convergence tests rather than causal interventions."
)

METHODS_HEADING = "Mechanistic pathway, attribution and counterfactual analyses"
METHODS_PARAS = [
    (
        "Mechanistic analyses were added after the primary claim-audited projection to ask which recurrent features carried the "
        "state-capacity distinction. In artificial agents, mediation models tested imposed perturbation family as the predictor, a "
        "single recurrent-dynamical feature as mediator and a behavioural fingerprint as outcome, with model family, seed and mean "
        "accuracy included as controls. Indirect effects were estimated as the product of standardised family-to-mediator and "
        "mediator-to-outcome paths and screened by bootstrap confidence intervals with false-discovery-rate correction."
    ),
    (
        "Feature-to-mechanism attribution used a grouped cross-validated random-forest classifier to distinguish state-like from "
        "capacity-like artificial perturbations from behavioural and dynamical fingerprint features. Groups were defined by architecture "
        "and seed to avoid evaluating on near-duplicates of the same trained recurrent model. Permutation importance on held-out folds "
        "quantified the contribution of each feature and was summarised by behavioural versus dynamical feature group."
    ),
    (
        "Counterfactual perturbation-response contrasts compared graph-preserving state perturbations with the matched baseline operating "
        "state at the same architecture, seed and hidden size, and compared graph-altering capacity scale with the h=1 capacity reference "
        "within architecture and seed. Human high-versus-low profile contrasts and human mediation models were treated as observational "
        "proxies only; they were used to test whether fitted recurrent geometry accounted for profile-behaviour covariance, not to infer "
        "experimental causality."
    ),
]


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    p.add_run(text)
    return p


def insert_picture_after(paragraph, image_path: Path, width_inches: float = 6.4):
    p = insert_paragraph_after(paragraph, "")
    p.alignment = 1
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def replace_text(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    paragraph.text = paragraph.text.replace(old, new)


def docx_integrity(path: Path) -> dict[str, int | str | None]:
    result: dict[str, int | str | None] = {}
    with zipfile.ZipFile(path) as zf:
        result["bad_member"] = zf.testzip()
        names = zf.namelist()
        result["duplicate_members"] = len(names) - len(set(names))
        for xml_name in ["[Content_Types].xml", "word/document.xml", "word/settings.xml", "word/_rels/document.xml.rels"]:
            etree.fromstring(zf.read(xml_name))
        result["media_count"] = len([n for n in names if n.startswith("word/media/")])
    d = Document(str(path))
    result["paragraph_count"] = len(d.paragraphs)
    result["inline_shapes"] = len(d.inline_shapes)
    return result


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    if not FIG.exists():
        raise FileNotFoundError(FIG)

    doc = Document(str(INPUT))

    # Abstract: add the mechanistic contribution without changing the cautious scalar-coordinate conclusion.
    abstract = doc.paragraphs[2]
    abstract.text = (
        "Cognitive science distinguishes a system's momentary operating state, such as arousal or response policy, from its more enduring capacity. "
        "A recent two-axis framework formalises this contrast for neural systems and predicts that state and capacity leave distinct fingerprints, but the two are confounded in human manipulations. "
        "We therefore built an external standard in recurrent neural networks, imposing graph-preserving interventions (gain, noise, decision temperature and threshold) and graph-altering interventions (workspace size and recurrent depth) before examining any human data. "
        "The two families were separable across architectures (balanced accuracy 0.92-0.94) and transferred to held-out networks. However, a hybrid-agent test rejected the stronger hypothesis that each axis is a recoverable scalar dial. "
        "Mechanistic analyses showed that the family boundary was carried primarily by recurrent-rank dynamics (cross-validated balanced accuracy 0.926; AUC 0.960), and that imposed artificial perturbation effects passed through recurrent rank, trajectory radius and hidden-step dynamics. "
        "Projecting 1,283 human observations from four open datasets onto the fixed profiles, capacity emerged as a convergent load-pressure and recurrent-geometry profile, whereas state was a within-person reliability profile whose scalar form remained unvalidated. "
        "Human observational pathway tests linked profile-behaviour covariance to fitted recurrent geometry, but were not interpreted as causal neural mediation. "
        "The findings support state and capacity as separable computational profiles while falsifying their strongest interpretation as two architecture-free scalar dials."
    )

    # Insert the new Results subsection after the Fig. 3 caption.
    anchor = doc.paragraphs[29]
    p = insert_paragraph_after(anchor, NEW_RESULTS_HEADING, style=doc.paragraphs[25].style.name)
    for text in NEW_RESULTS_PARAS:
        p = insert_paragraph_after(p, text)
    p = insert_picture_after(p, FIG)
    p = insert_paragraph_after(p, NEW_FIG_CAPTION)

    # Renumber downstream figure captions.
    for para in doc.paragraphs:
        if para.text.startswith("Fig. 6 |"):
            para.text = para.text.replace("Fig. 6 |", "Fig. 7 |", 1)
        elif para.text.startswith("Fig. 5 |"):
            para.text = para.text.replace("Fig. 5 |", "Fig. 6 |", 1)
        elif para.text.startswith("Fig. 4 | Capacity"):
            para.text = para.text.replace("Fig. 4 |", "Fig. 5 |", 1)

    # Discussion updates.
    for para in doc.paragraphs:
        if para.text.startswith("The study resolves the state-capacity hypothesis at two levels."):
            para.text = (
                "The study resolves the state-capacity hypothesis at three levels. At the level of kinds, the distinction is supported: "
                "graph-preserving and graph-altering perturbations leave separable recurrent-network signatures, and human behaviour partly aligns "
                "with the resulting profiles. At the mechanistic level, the family boundary is carried chiefly by recurrent state-space organisation: "
                "rank, trajectory radius and hidden-step dynamics mediate artificial perturbation effects and account for part of the human profile-behaviour "
                "covariance. At the level of scalar magnitudes, the strong hypothesis is rejected: after residualising architecture and the complementary "
                "intervention family, neither state nor capacity is validated as a general architecture-free scalar coordinate. The traditional distinction "
                "between operating state and structural capacity is not wrong, but it is better represented as separable computational profiles than as two clean cognitive dials."
            )
        elif para.text.startswith("Capacity is the stronger human construct in the current evidence."):
            para.text = (
                "Capacity is the stronger human construct in the current evidence. Its profile is expressed in load pressure, recurrent-trajectory geometry, "
                "directional generalisation and mechanistic pathway analyses. This does not make it general intelligence, and it does not supersede psychometric "
                "factors3,40-43. Instead, the network standard adds a mechanistic anchor: capacity-related profiles are tied to representational-resource "
                "manipulations in artificial systems, to recurrent-rank and trajectory-geometry changes in those systems, and to the dimensionality and "
                "load-sensitivity of fitted recurrent dynamics in humans. That combination is more specific than a purely factor-analytic summary but less "
                "sweeping than a claim about a single universal capacity score."
            )
        elif para.text.startswith("State is informative in a different way."):
            para.text = (
                "State is informative in a different way. Its lower ICC and trial-count dependence are exactly what one should expect from a within-person "
                "operating profile. Lapses, drift, response variability and reliability can move together, but their mapping to a single scalar severity is "
                "not stable across architectures or datasets. The mechanistic screen reinforces this interpretation: graph-preserving state interventions "
                "primarily changed response entropy, hidden variability and probability volatility rather than expanding recurrent rank. State should therefore "
                "be treated as a dense-data profile of current operating reliability and instability. This interpretation also explains why raw scalar recovery "
                "can look strong in artificial agents: raw fingerprints retain architecture and shared family variance. Residualisation removes that shared variance, "
                "and the scalar gate appropriately fails."
            )
        elif para.text.startswith("Several limitations are important."):
            para.text = (
                "Several limitations are important. First, the artificial standard is restricted to recurrent architectures; attention-based, transformer-like and "
                "multimodal architectures remain untested, so capacity is probed here along the depth and recurrence dimensions of the framework but not its breadth, "
                "or multimodal-binding, dimension. Second, capacity is participant-level in the current implementation and partly includes load-sensitive components, "
                "so load-pressure evidence should be read as convergent profile evidence, not as an independent causal manipulation of capacity. Third, the load-pressure "
                "display uses observational participant differences under a task manipulation rather than an experimental manipulation of structural resource. Fourth, "
                "the human pathway analyses are observational and should be read as recurrent-dynamics convergence, not as causal mediation. Fifth, the physiology screen "
                "uses summary features and is exploratory. Sixth, ds007554 contains reconstructed correctness for part of the record. Finally, the work was not prospectively "
                "preregistered; gates and the claim audit were fixed in the analysis registry before final projection and reporting, but the result should be treated as a "
                "claim-audited secondary analysis of open datasets."
            )

    # Insert Methods section before Capacity load-pressure models.
    methods_anchor_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text == "Capacity load-pressure models":
            methods_anchor_idx = idx
            break
    if methods_anchor_idx is None:
        raise RuntimeError("Could not locate Methods insertion point")
    methods_anchor = doc.paragraphs[methods_anchor_idx]
    p = insert_paragraph_after(
        doc.paragraphs[methods_anchor_idx - 1],
        METHODS_HEADING,
        style=methods_anchor.style.name,
    )
    for text in METHODS_PARAS:
        p = insert_paragraph_after(p, text)

    # Update final extended-data sentence.
    for para in doc.paragraphs:
        if para.text.startswith("Extended Data Fig. 1 and Extended Data Tables"):
            para.text = (
                "Extended Data Fig. 1 and Extended Data Tables 1-5 provide the model-comparison baseline, artificial architecture and hybrid-recovery tables, "
                "mechanistic pathway and counterfactual source data, load-pressure coefficients, state variance components, capacity ablation/variant checks, "
                "physiology associations and the full claim-audit record. These tables are essential for reviewers because the main text deliberately distinguishes "
                "profile-level support from failed scalar-coordinate gates."
            )

    doc.save(OUTPUT)
    integrity = docx_integrity(OUTPUT)
    CHANGE_LOG.write_text(
        "\n".join(
            [
                "# Manuscript mechanistic revision log",
                "",
                f"Input: `{INPUT}`",
                f"Output: `{OUTPUT}`",
                "",
                "Inserted a new Results subsection after Fig. 3 and before the capacity-profile section.",
                "Inserted the mechanistic insights figure as new Fig. 4 and renumbered previous Figs. 4-6 to Figs. 5-7.",
                "Updated Abstract, Discussion, Methods and Extended Data summary language.",
                "",
                "Key added claims:",
                "- Artificial perturbation family effects pass through recurrent rank, trajectory radius and hidden-step dynamics.",
                "- Recurrent-rank dynamics dominates feature attribution for the state-capacity family boundary.",
                "- Human pathway tests are observational convergence through fitted recurrent geometry, not causal neural mediation.",
                "",
                f"Integrity check: {integrity}",
            ]
        ),
        encoding="utf-8",
    )
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {CHANGE_LOG}")
    print(integrity)


if __name__ == "__main__":
    main()
